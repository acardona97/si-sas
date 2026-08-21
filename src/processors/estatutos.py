# -*- coding: utf-8 -*-
"""
Generador de Estatutos de Constitución S.A.S. (.docx)

Usa python-docx para editar la plantilla estatutos_template.docx
reemplazando los tokens {{TOKEN}} con los datos del formulario.
"""
from copy import deepcopy
from datetime import date
from docx import Document
from docx.shared import RGBColor, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


# ════════════════════════════════════════════════════════════════
# UTILIDADES
# ════════════════════════════════════════════════════════════════

UNIDADES = [
    "", "un", "dos", "tres", "cuatro", "cinco", "seis", "siete",
    "ocho", "nueve", "diez", "once", "doce", "trece", "catorce",
    "quince", "dieciséis", "diecisiete", "dieciocho", "diecinueve",
]
DECENAS = [
    "", "", "veinte", "treinta", "cuarenta", "cincuenta",
    "sesenta", "setenta", "ochenta", "noventa",
]
CENTENAS = [
    "", "cien", "doscientos", "trescientos", "cuatrocientos",
    "quinientos", "seiscientos", "setecientos", "ochocientos", "novecientos",
]

MESES = [
    "", "enero", "febrero", "marzo", "abril", "mayo", "junio",
    "julio", "agosto", "septiembre", "octubre", "noviembre", "diciembre",
]


def _numero_a_letras(n):
    """Convierte un número entero a su representación en letras (español)."""
    if n == 0:
        return "cero"
    if n < 0:
        return "menos " + _numero_a_letras(-n)

    resultado = ""

    if n >= 1_000_000_000:
        miles_m = n // 1_000_000_000
        if miles_m == 1:
            resultado += "mil millones"
        else:
            resultado += _numero_a_letras(miles_m) + " mil millones"
        n %= 1_000_000_000
        if n > 0:
            resultado += " "

    if n >= 1_000_000:
        millones = n // 1_000_000
        if millones == 1:
            resultado += "un millón"
        else:
            resultado += _numero_a_letras(millones) + " millones"
        n %= 1_000_000
        if n > 0:
            resultado += " "

    if n >= 1000:
        miles = n // 1000
        if miles == 1:
            resultado += "mil"
        else:
            resultado += _numero_a_letras(miles) + " mil"
        n %= 1000
        if n > 0:
            resultado += " "

    if n >= 100:
        c = n // 100
        if c == 1 and n % 100 > 0:
            resultado += "ciento "
        else:
            resultado += CENTENAS[c]
            if n % 100 > 0:
                resultado += " "
        n %= 100

    if n >= 20:
        d = n // 10
        u = n % 10
        if d == 2 and u > 0:
            # 21-29 con acentos correctos en español
            veinti_map = {
                1: "veintiún",
                2: "veintidós",
                3: "veintitrés",
                4: "veinticuatro",
                5: "veinticinco",
                6: "veintiséis",
                7: "veintisiete",
                8: "veintiocho",
                9: "veintinueve",
            }
            resultado += veinti_map[u]
        else:
            resultado += DECENAS[d]
            if u > 0:
                resultado += " y " + UNIDADES[u]
    elif n > 0:
        resultado += UNIDADES[n]

    return resultado.strip()


def _monto_letras_cifras(n):
    """Ej: 1000000 -> 'UN MILLÓN DE PESOS (COP $1.000.000)'.

    Regla gramatical: solo se usa "DE PESOS" cuando el número termina en
    millón/millones/billón/billones (y nada más después). En otros casos
    ("seiscientos mil", "tres millones quinientos mil", "doscientos mil")
    se usa simplemente "PESOS".
    """
    letras = _numero_a_letras(n).upper()
    cifras = f"{n:,}".replace(",", ".")
    # ¿Termina exactamente en "MILLÓN" o "MILLONES" / "BILLÓN" o "BILLONES"?
    palabras = letras.split()
    if palabras and palabras[-1] in ("MILLÓN", "MILLONES", "BILLÓN", "BILLONES"):
        conector = "DE PESOS"
    else:
        conector = "PESOS"
    return f"{letras} {conector} (COP ${cifras})"


def _acciones_letras_cifras(n):
    """Ej: 1000000 -> 'un millón (1.000.000)'."""
    letras = _numero_a_letras(n)
    cifras = f"{n:,}".replace(",", ".")
    return f"{letras} ({cifras})"


def _valor_nominal_frase(valor, sufijo="moneda legal colombiana"):
    """Frase del valor nominal por acción, concordada en número.

    Ej: 1    -> 'un peso moneda legal colombiana (COP $1)'
        100  -> 'cien pesos moneda legal colombiana (COP $100)'

    El sufijo varía entre artículos de la plantilla ('moneda legal
    colombiana' en los artículos 4 y 5, 'colombiano/s' en el 6), por lo que
    se recibe como parámetro.
    """
    letras = _numero_a_letras(valor)
    cifras = f"{valor:,}".replace(",", ".")
    if valor == 1:
        sustantivo = "peso"
        suf = sufijo.replace("colombianos", "colombiano")
    else:
        sustantivo = "pesos"
        suf = sufijo.replace("colombiano", "colombianos")
    return f"{letras} {sustantivo} {suf} (COP ${cifras})"


def _fecha_literal(f):
    """Ej: date(2026,5,20) -> 'veinte (20) días del mes de mayo de 2026'."""
    dia_letras = _numero_a_letras(f.day)
    return f"{dia_letras} ({f.day}) días del mes de {MESES[f.month]} de {f.year}"


def _genero(g, masc, fem):
    """Conjuga según género: 'M' -> masc, 'F' -> fem."""
    return fem if str(g).upper() == "F" else masc


# ════════════════════════════════════════════════════════════════
# DETECCIÓN AUTOMÁTICA DE GÉNERO POR NOMBRE
# ════════════════════════════════════════════════════════════════
# Base de datos de nombres propios comunes en Colombia (sin acentos).
# Se compara contra el PRIMER nombre del campo "nombre" del usuario, por lo
# que detecta correctamente compuestos como "María Camila" (→F) o "Juan
# Pablo" (→M).
#
# Si la detección automática es exitosa, prevalece sobre el dropdown del
# formulario (para evitar errores como marcar a "Yaneth" como masculino
# por dejar el default sin cambiar).

NOMBRES_FEMENINOS = frozenset({
    # A
    "abigail", "adela", "adriana", "alba", "alejandra", "alexandra", "alicia",
    "amalia", "amelia", "ana", "andrea", "angela", "angelica", "antonia",
    "ariana", "arianna", "astrid", "aura",
    # B
    "beatriz", "berenice", "betty", "blanca", "brenda", "briana",
    # C
    "camila", "carla", "carmen", "carolina", "catalina", "cecilia", "celia",
    "cindy", "claudia", "constanza", "cristina",
    # D
    "daniela", "danna", "delfina", "diana", "dolores", "dora",
    # E
    "edith", "elena", "eleonora", "elisa", "elizabeth", "elsa", "elvira",
    "emma", "erika", "esperanza", "estefania", "estela", "estrella", "eugenia",
    "eva",
    # F
    "fabiana", "fanny", "felipa", "fernanda", "florencia", "francisca", "frida",
    # G
    "gabriela", "geraldine", "gertrudis", "gladys", "gloria", "graciela",
    "guadalupe",
    # H
    "helena", "hilda",
    # I
    "ileana", "ines", "ingrid", "irene", "isabel", "isabella", "isadora",
    "ivana", "ivonne",
    # J
    "jacqueline", "jazmin", "jeimy", "jenifer", "jennifer", "jenny", "jessica",
    "jimena", "joanna", "johana", "johanna", "josefa", "josefina", "judith",
    "julia", "juliana", "july",
    # K
    "karen", "karina", "karla", "katherine", "katia", "kelly",
    # L
    "laura", "leticia", "liliana", "lina", "lisa", "lola", "lorena",
    "lourdes", "lucia", "luisa", "luz",
    # M
    "magdalena", "manuela", "marcela", "margarita", "maria", "mariana",
    "maribel", "maritza", "marlene", "marta", "martha", "matilda", "melissa",
    "mercedes", "milena", "monica", "myriam", "miryam",
    # N
    "nadia", "nancy", "natalia", "natasha", "nayeli", "nelly", "nicol",
    "nicole", "nidia", "nieves", "norma", "nubia",
    # O
    "olga", "olivia", "ortencia",
    # P
    "paloma", "paola", "patricia", "paulina", "pilar",
    # R
    "rafaela", "raquel", "rebeca", "regina", "rocio", "rosa", "rosalia",
    "rosario", "ruth",
    # S
    "sandra", "sara", "sarai", "silvia", "sofia", "soledad", "sonia", "stella",
    "susana",
    # T
    "tania", "tatiana", "teresa", "trinidad",
    # V
    "valentina", "valeria", "vanessa", "veronica", "victoria", "viviana",
    # W
    "wendy",
    # X
    "ximena",
    # Y
    "yadira", "yaneth", "yanira", "yazmin", "yenny", "yesenia", "yolanda",
    "yuliana", "yuri",
    # Z
    "zaida", "zoila",
})

NOMBRES_MASCULINOS = frozenset({
    # A
    "abel", "abraham", "adolfo", "adrian", "adriano", "agustin", "alan",
    "alberto", "alejandro", "alex", "alexander", "alfonso", "alfredo", "alvaro",
    "anderson", "andres", "angel", "anibal", "antonio", "arnoldo", "arturo",
    "augusto",
    # B
    "baltasar", "benjamin", "bernardo", "boris", "brandon", "brayan",
    # C
    "camilo", "carlos", "cesar", "cristian", "cristobal", "cristopher",
    # D
    "daniel", "dario", "david", "diego", "dilan", "domingo",
    # E
    "eduardo", "edwin", "elias", "emilio", "emmanuel", "enrique", "erick",
    "ernesto", "esteban", "eugenio",
    # F
    "fabian", "fabio", "federico", "felipe", "fernando", "francisco", "frank",
    # G
    "gabriel", "gerardo", "german", "gilberto", "gonzalo", "guido", "guillermo",
    "gustavo",
    # H
    "hector", "henry", "heriberto", "hernan", "hernando", "hugo", "humberto",
    # I
    "ivan", "ignacio", "isidro", "ismael", "isaac",
    # J
    "jacobo", "jaime", "javier", "jefferson", "jeronimo", "jesus", "jhon",
    "joaquin", "joel", "john", "jonathan", "jorge", "jose", "josue", "juan",
    "julian", "julio",
    # K
    "kevin",
    # L
    "leonardo", "leopoldo", "luis", "luciano",
    # M
    "manuel", "marco", "marcos", "mario", "mateo", "matias", "mauricio",
    "miguel",
    # N
    "nelson", "nestor", "nicolas", "norberto",
    # O
    "octavio", "omar", "orlando", "oscar", "osvaldo", "otoniel",
    # P
    "pablo", "patricio", "pedro", "phillip",
    # R
    "rafael", "ramiro", "ramon", "raul", "ricardo", "roberto", "rodrigo",
    "rolando", "ronald", "ruben",
    # S
    "salvador", "samuel", "santiago", "saul", "sebastian", "sergio", "simon",
    # T
    "tadeo", "tomas", "tulio",
    # V
    "valentin", "vicente", "victor",
    # W
    "wilfredo", "william", "wilson",
    # Y
    "yair",
    # Z
    "zacarias",
})


def _strip_acentos(texto):
    """Quita acentos de un texto. 'María' → 'maria'."""
    import unicodedata
    return "".join(
        c for c in unicodedata.normalize("NFKD", texto.lower())
        if not unicodedata.combining(c)
    )


def _detect_gender_from_name(nombre):
    """
    Detecta el género de una persona a partir de su PRIMER nombre.

    Retorna 'F', 'M' o None (si no se puede determinar).

    Ejemplos:
      "Andrea López Ríos"   → 'F'
      "María Camila Torres" → 'F'
      "Juan Pablo García"   → 'M'
      "Madonna"             → None (no está en la lista)
      "J. P. González"      → None (iniciales)
    """
    if not nombre:
        return None
    palabras = nombre.strip().split()
    if not palabras:
        return None
    primer = palabras[0]
    # Saltar iniciales (ej. "J.", "M.A.")
    if len(primer) < 2 or "." in primer:
        return None
    primer_norm = _strip_acentos(primer)
    if primer_norm in NOMBRES_FEMENINOS:
        return "F"
    if primer_norm in NOMBRES_MASCULINOS:
        return "M"
    return None


def _resolve_gender(nombre, genero_usuario="M"):
    """
    Determina el género combinando detección automática del nombre con la
    selección manual del usuario.

    Política: la detección automática prevalece sobre la selección manual
    cuando puede identificar el nombre con certeza. Esto evita errores
    comunes como dejar "Yaneth" marcada como masculino por olvidar cambiar
    el dropdown.
    """
    auto = _detect_gender_from_name(nombre)
    if auto is not None:
        return auto
    return (genero_usuario or "M").upper()


# ════════════════════════════════════════════════════════════════
# BLOQUES DINÁMICOS
# ════════════════════════════════════════════════════════════════

def _bloque_comparecencia(accionistas):
    """Construye el bloque de comparecencia de accionistas."""
    lineas = []
    for i, acc in enumerate(accionistas, 1):
        tipo = acc.get("tipo", "natural")
        if tipo == "juridica":
            rl_genero = acc.get("rl_genero", "M")
            ident_rl = _genero(rl_genero, "identificado", "identificada")
            linea = (
                f"{i}. {acc['nombre'].upper()}, sociedad comercial identificada "
                f"con NIT {acc.get('id_num', '')}, domiciliada en "
                f"{acc.get('domicilio', 'Medellín')}, representada legalmente por "
                f"{acc.get('rl_nombre', '').upper()}, mayor de edad, "
                f"{ident_rl} con C.C. No. {acc.get('rl_cc', '')}, "
                f"expedida en {acc.get('rl_expedicion', '')}."
            )
        else:
            genero = acc.get("genero", "M")
            id_tipo = acc.get("id_tipo", "C.C.")
            linea = (
                f"{i}. {acc['nombre'].upper()}, mayor de edad, "
                f"{_genero(genero, 'identificado', 'identificada')} con "
                f"{id_tipo} No. {acc.get('id_num', '')}, expedida en "
                f"{acc.get('expedicion', '')}, "
                f"{_genero(genero, 'domiciliado', 'domiciliada')} en "
                f"{acc.get('domicilio', '')}."
            )
        lineas.append(linea)
    return "\n".join(lineas)


def _bloque_firmas(accionistas, rl_principal, nombre_sas):
    """Construye el bloque de firmas finales."""
    lineas = []

    if len(accionistas) == 1:
        acc = accionistas[0]
        nombre = acc["nombre"].upper()
        id_num = acc.get("id_num", "")
        id_tipo = acc.get("id_tipo", "C.C.")

        rl_nombre = rl_principal.get("nombre", "").upper()

        if nombre == rl_nombre:
            # Caso 1: único accionista que también es RL
            lineas.append(nombre)
            lineas.append(f"{id_tipo} No. {id_num}")
            lineas.append(
                f"En calidad de constituyente único y representante legal "
                f"principal de {nombre_sas}"
            )
        else:
            # Caso 2: único accionista distinto del RL
            lineas.append(nombre)
            lineas.append(f"{id_tipo} No. {id_num}")
            lineas.append(f"En calidad de constituyente único de {nombre_sas}")
            lineas.append("")
            lineas.append(rl_nombre)
            lineas.append(f"C.C. No. {rl_principal.get('cc', '')}")
            lineas.append(
                f"En calidad de representante legal principal de {nombre_sas}"
            )
    else:
        # Caso 3: múltiples accionistas
        for acc in accionistas:
            tipo = acc.get("tipo", "natural")
            if tipo == "juridica":
                lineas.append(acc.get("rl_nombre", "").upper())
                lineas.append(f"C.C. No. {acc.get('rl_cc', '')}")
                lineas.append(
                    f"En calidad de representante legal de {acc['nombre'].upper()}, "
                    f"constituyente de {nombre_sas}"
                )
            else:
                lineas.append(acc["nombre"].upper())
                lineas.append(f"{acc.get('id_tipo', 'C.C.')} No. {acc.get('id_num', '')}")
                lineas.append(f"En calidad de constituyente de {nombre_sas}")
            lineas.append("")

        # Si el RL no es ninguno de los accionistas, agregar
        rl_nombre = rl_principal.get("nombre", "").upper()
        acc_nombres = [a["nombre"].upper() for a in accionistas]
        rl_nombres_pj = [a.get("rl_nombre", "").upper() for a in accionistas if a.get("tipo") == "juridica"]
        if rl_nombre not in acc_nombres and rl_nombre not in rl_nombres_pj:
            lineas.append(rl_nombre)
            lineas.append(f"C.C. No. {rl_principal.get('cc', '')}")
            lineas.append(
                f"En calidad de representante legal principal de {nombre_sas}"
            )

    return "\n".join(lineas)


# ════════════════════════════════════════════════════════════════
# TABLA DE ACCIONISTAS
# ════════════════════════════════════════════════════════════════

def _fmt_money_table(n):
    """Formatea número como moneda: 1000000 -> '$1.000.000'."""
    return f"${n:,}".replace(",", ".")


def _fmt_num_table(n):
    """Formatea número con puntos: 1000000 -> '1.000.000'."""
    return f"{n:,}".replace(",", ".")


def _set_cell_text(cell, text, bold=False):
    """Reemplaza el texto de una celda preservando el formato de los runs."""
    para = cell.paragraphs[0]
    if para.runs:
        para.runs[0].text = text
        if bold:
            para.runs[0].bold = True
        for run in para.runs[1:]:
            run.text = ""
    else:
        run = para.add_run(text)
        run.font.name = "Cambria"
        if bold:
            run.bold = True


def _fill_accionistas_table(doc, accionistas, capital_suscrito, capital_pagado,
                            valor_nominal=1):
    """
    Llena la tabla de accionistas/acciones/capital en los estatutos.

    La tabla (doc.tables[0]) tiene:
      - Row 0: Encabezados (Accionista, Identificación, %, No. Acciones, Capital Suscrito, Capital Pagado)
      - Row 1-2: Filas placeholder con guiones bajos
      - Row 3: Fila TOTAL

    Se clonan las filas existentes para preservar bordes, fuentes y alineación.
    """
    # Se localiza por el encabezado y no por índice: el Artículo Primero
    # Transitorio puede insertar antes la tabla de la junta directiva, con lo
    # que la de accionistas deja de ser doc.tables[0].
    table = None
    for t in doc.tables:
        if t.rows and "accionista" in t.rows[0].cells[0].text.strip().lower():
            table = t
            break
    if table is None:
        return

    tbl = table._tbl

    # Guardar templates (deep copy del XML) antes de eliminar
    data_tr_template = deepcopy(table.rows[1]._tr)
    total_tr_template = deepcopy(table.rows[3]._tr)

    # Eliminar filas 1, 2, 3 (solo conservar encabezado row 0)
    for row in list(table.rows)[1:]:
        tbl.remove(row._tr)

    # Calcular datos por accionista y agregar filas
    total_acciones = 0
    total_suscrito = 0
    total_pagado = 0

    vn = max(1, int(valor_nominal or 1))

    for acc in accionistas:
        pct = float(acc.get("porcentaje", 0))
        suscrito_acc = int(capital_suscrito * pct / 100)
        acciones = suscrito_acc // vn
        # Capital pagado individual: puede ser distinto por accionista
        # (algunos pagan todo, otros parte). Si no se especificó, se
        # calcula proporcional al global.
        if "capital_pagado_num" in acc:
            pagado_acc = acc["capital_pagado_num"]
        else:
            pagado_acc = int(capital_pagado * pct / 100)

        total_acciones += acciones
        total_suscrito += suscrito_acc
        total_pagado += pagado_acc

        # Clonar fila template y agregar al final de la tabla
        new_tr = deepcopy(data_tr_template)
        tbl.append(new_tr)
        new_row = table.rows[-1]

        # Datos del accionista
        nombre = acc.get("nombre", "").upper()
        tipo = acc.get("tipo", "natural")
        if tipo == "juridica":
            id_label = f"NIT {acc.get('id_num', '')}"
        else:
            id_label = f"{acc.get('id_tipo', 'C.C.')} {acc.get('id_num', '')}"

        _set_cell_text(new_row.cells[0], nombre, bold=True)   # nombre propio
        _set_cell_text(new_row.cells[1], id_label)
        _set_cell_text(new_row.cells[2], f"{int(pct)}%")
        _set_cell_text(new_row.cells[3], _fmt_num_table(acciones))
        _set_cell_text(new_row.cells[4], _fmt_money_table(suscrito_acc))
        _set_cell_text(new_row.cells[5], _fmt_money_table(pagado_acc))

    # Agregar fila TOTAL
    new_total_tr = deepcopy(total_tr_template)
    tbl.append(new_total_tr)
    total_row = table.rows[-1]

    _set_cell_text(total_row.cells[0], "TOTAL")
    _set_cell_text(total_row.cells[1], "")
    _set_cell_text(total_row.cells[2], "100%")
    _set_cell_text(total_row.cells[3], _fmt_num_table(total_acciones))
    _set_cell_text(total_row.cells[4], _fmt_money_table(total_suscrito))
    _set_cell_text(total_row.cells[5], _fmt_money_table(total_pagado))

    # La plantilla trae las filas de encabezado y total en azul claro; se
    # repintan en gris para que ambos cuadros del documento coincidan.
    _pintar_fila(table.rows[0], TBL_RELLENO_ENCABEZADO)
    _pintar_fila(total_row, TBL_RELLENO_ENCABEZADO)


# ════════════════════════════════════════════════════════════════
# POST-PROCESAMIENTO ESTÉTICO
# ════════════════════════════════════════════════════════════════

def _split_comparecencia_paragraph(doc):
    """
    Convierte todo salto de línea (<w:br/>) en un párrafo independiente.

    Word justifica la línea que antecede a un <w:br/> como si fuera una línea
    intermedia, estirando los espacios de punta a punta. En un texto
    justificado eso deja huecos enormes entre las últimas palabras. La única
    forma de evitarlo es que cada bloque sea un párrafo de verdad, no una
    línea partida dentro del mismo párrafo.

    Aplica tanto a la comparecencia de accionistas como al objeto social, que
    llega redactado en varios párrafos, y a cualquier otro token multilínea.
    Cada párrafo nuevo se clona del original, así conserva alineación,
    interlineado, sangría y fuente.
    """
    body = doc._element.body

    for p_elem in list(body.findall(qn("w:p"))):
        if p_elem.find(f".//{qn('w:br')}") is None:
            continue

        main_run = None
        for r in p_elem.findall(qn("w:r")):
            if r.find(qn("w:br")) is not None:
                main_run = r
                break
        if main_run is None:
            continue

        # Trocear el contenido del run por los saltos
        partes, actual = [], ""
        for child in list(main_run):
            tag = child.tag.split("}")[-1]
            if tag == "t":
                actual += child.text or ""
            elif tag == "br":
                partes.append(actual)
                actual = ""
        partes.append(actual)

        # Un tramo vacío es la línea en blanco que separa dos párrafos y se
        # emite como párrafo vacío, que es lo que da el espaciado del
        # documento. Solo se descartan los de los extremos.
        while partes and not partes[0].strip():
            partes.pop(0)
        while partes and not partes[-1].strip():
            partes.pop()
        if not partes:
            continue

        # El run original se queda con el primer tramo, ya sin saltos
        for child in list(main_run):
            if child.tag.split("}")[-1] in ("t", "br"):
                main_run.remove(child)
        t = OxmlElement("w:t")
        t.text = partes[0]
        t.set(qn("xml:space"), "preserve")
        main_run.append(t)

        # Plantilla de run para los tramos siguientes: se copia el formato del
        # run que traía los saltos, no el del párrafo entero. Si el párrafo
        # empieza con una etiqueta en negrilla ("Objeto social: "), esa
        # etiqueta pertenece solo al primer tramo y no debe repetirse.
        run_modelo = deepcopy(main_run)
        for child in list(run_modelo):
            if child.tag.split("}")[-1] in ("t", "br"):
                run_modelo.remove(child)

        insert_after = p_elem
        for parte in partes[1:]:
            clon = deepcopy(p_elem)
            # El clon se queda solo con las propiedades del párrafo
            for r in clon.findall(qn("w:r")):
                clon.remove(r)

            # Un tramo de continuación no es un elemento nuevo de la lista: si
            # conservara la numeración automática se comería el número del
            # artículo siguiente ("Artículo 4." aparecería a mitad del objeto
            # social y el capital autorizado pasaría a ser el 5).
            clon_pPr = clon.find(qn("w:pPr"))
            if clon_pPr is not None:
                for numPr in clon_pPr.findall(qn("w:numPr")):
                    clon_pPr.remove(numPr)

            nuevo_run = deepcopy(run_modelo)
            t_new = OxmlElement("w:t")
            t_new.text = parte
            t_new.set(qn("xml:space"), "preserve")
            nuevo_run.append(t_new)
            clon.append(nuevo_run)

            insert_after.addnext(clon)
            insert_after = clon


def _reemplazar_frase(doc, viejo, nuevo):
    """
    Sustituye una frase literal en el cuerpo, aunque el párrafo no traiga
    tokens.

    `_replace_in_paragraph` solo actúa sobre párrafos con "{{", así que los
    textos fijos de la plantilla —como la frase del artículo de nombramiento
    del representante legal— necesitan este camino. Se reemplaza dentro del
    run que la contiene, de modo que el resto del párrafo conserva su formato
    (por ejemplo el título en negrilla que lo antecede).
    """
    if not viejo or not nuevo:
        return False
    for p_elem in doc._element.body.findall(qn("w:p")):
        if viejo not in _get_para_text(p_elem):
            continue
        for run in p_elem.findall(qn("w:r")):
            t_elems = run.findall(qn("w:t"))
            if len(t_elems) != 1:
                continue
            contenido = t_elems[0].text or ""
            if viejo in contenido:
                t_elems[0].text = contenido.replace(viejo, nuevo)
                t_elems[0].set(qn("xml:space"), "preserve")
                return True
    return False


def _colapsar_espacios_dobles(doc):
    """
    Deja un solo espacio entre palabras en todo el documento.

    Los espacios de más vienen de dos lados: de la plantilla, que en algunos
    puntos tiene dos espacios seguidos alrededor de un token, y de los valores
    sustituidos que terminan o empiezan con espacio. En texto justificado se
    notan como huecos.

    Recorre los <w:t> de cada párrafo en orden, de modo que también detecta el
    espacio doble repartido entre dos runs distintos.
    """
    for p_elem in doc._element.body.iter(qn("w:p")):
        anterior_es_espacio = False
        for t in p_elem.findall(f".//{qn('w:t')}"):
            texto = t.text or ""
            if not texto:
                continue
            salida = []
            for ch in texto:
                if ch == " ":
                    if anterior_es_espacio:
                        continue
                    anterior_es_espacio = True
                else:
                    anterior_es_espacio = False
                salida.append(ch)
            nuevo = "".join(salida)
            if nuevo != texto:
                t.text = nuevo
                t.set(qn("xml:space"), "preserve")


def _enable_auto_hyphenation(doc):
    """
    Habilita silabeo automático en el documento.

    En texto justificado (jc="both"), Word distribuye el espacio sobrante
    entre las palabras de cada línea.  Cuando una palabra larga no cabe al
    final, la línea queda con pocas palabras y espacios muy anchos.

    El silabeo automático permite que Word corte palabras en sílabas,
    llenando mejor cada línea y reduciendo los huecos antiestéticos.
    """
    settings = doc.settings._element
    # <w:autoHyphenation w:val="true"/>
    auto_hyph = settings.find(qn("w:autoHyphenation"))
    if auto_hyph is None:
        auto_hyph = OxmlElement("w:autoHyphenation")
        settings.append(auto_hyph)
    auto_hyph.set(qn("w:val"), "true")


def _post_process_doc(doc, nombre_sas):
    """
    Correcciones estéticas después de reemplazar tokens:
    1. Título de la sociedad: color negro (no rojo)
    2. Artículos numerados: eliminar highlight amarillo del template
    3. Espaciado: eliminar párrafos vacíos extra tras la comparecencia
    """
    for para in doc.paragraphs:
        text = para.text.strip()

        # ── 1. Título de la sociedad: quitar color rojo ──
        if text == nombre_sas:
            for run in para.runs:
                if run.font.color and run.font.color.rgb == RGBColor(0xEE, 0x00, 0x00):
                    run.font.color.rgb = RGBColor(0, 0, 0)

        # ── 2. Eliminar highlight amarillo de párrafos con numeración ──
        pPr = para._element.find(qn("w:pPr"))
        if pPr is not None:
            rPr = pPr.find(qn("w:rPr"))
            if rPr is not None:
                hl = rPr.find(qn("w:highlight"))
                if hl is not None:
                    rPr.remove(hl)
        # También quitar highlight de los runs individuales
        for run in para.runs:
            rPr = run._element.find(qn("w:rPr"))
            if rPr is not None:
                hl = rPr.find(qn("w:highlight"))
                if hl is not None:
                    rPr.remove(hl)

    # ── 3. Eliminar párrafos vacíos extra tras la comparecencia ──
    _remove_extra_empty_paras(doc)


def _remove_extra_empty_paras(doc):
    """Elimina párrafos vacíos consecutivos (máximo 1 vacío seguido)."""
    body = doc._element.body
    prev_empty = False
    to_remove = []

    for p_elem in body.findall(qn("w:p")):
        # Extraer texto del párrafo
        texts = [t.text or "" for t in p_elem.findall(f".//{qn('w:t')}")]
        text = "".join(texts).strip()

        if text == "":
            if prev_empty:
                to_remove.append(p_elem)
            prev_empty = True
        else:
            prev_empty = False

    for p_elem in to_remove:
        parent = p_elem.getparent()
        if parent is not None:
            parent.remove(p_elem)


def _insert_firmas_paragraphs(doc, accionistas, rl_principal, rl_suplente, nombre_sas,
                              rl_principales=None, rl_suplentes=None):
    """
    Reemplaza el párrafo de firmas (vacío tras token replacement) con
    bloques de firma bien espaciados y alineados a la izquierda.

    Cada firmante tiene:
      - 3 líneas vacías para espacio de firma manuscrita
      - Nombre (negrita)
      - Tipo y número de documento
      - Calidad (constituyente / representante legal)
    """
    # Encontrar el párrafo vacío donde estaba {{BLOQUE_FIRMAS_FINALES}}
    # (ahora es "" porque se reemplazó con cadena vacía)
    target_elem = None
    poderdante_elem = None  # Para eliminarlo al final (siempre sobra)
    body = doc._element.body
    all_paras = body.findall(qn("w:p"))

    # Buscar entre los últimos 30 párrafos el que tenía las firmas
    # (estará vacío, justo antes de otros párrafos vacíos al final)
    for p_elem in reversed(all_paras[-40:]):
        texts = [t.text or "" for t in p_elem.findall(f".//{qn('w:t')}")]
        text = "".join(texts).strip()
        if text == "En calidad de poderdante":
            # Encontramos el marcador de referencia — las firmas van ANTES
            target_elem = p_elem
            poderdante_elem = p_elem
            break

    if target_elem is None:
        # Fallback: buscar último párrafo con contenido y usar el siguiente vacío
        for i, p_elem in enumerate(all_paras):
            texts = [t.text or "" for t in p_elem.findall(f".//{qn('w:t')}")]
            text = "".join(texts).strip()
            if "proceden a firmarlo" in text:
                # Las firmas van después de este párrafo
                target_elem = all_paras[i + 1] if i + 1 < len(all_paras) else None
                break

    if target_elem is None:
        return

    # Construir lista de firmantes
    firmantes = _build_firmantes_list(accionistas, rl_principal, rl_suplente,
                                      nombre_sas, rl_principales, rl_suplentes)

    # Insertar párrafos de firma antes del target (o después del "proceden a firmarlo")
    insert_point = target_elem
    for firmante in firmantes:
        # Un bloque de firma es indivisible: espacio para la rúbrica, nombre,
        # documento, calidad y, si aplica, la aceptación del cargo. Se arma
        # completo y se marca para que Word no lo parta entre dos páginas.
        lineas = [""] * 4                       # espacio para firma manuscrita
        lineas.append(firmante["nombre"])
        lineas.append(firmante["documento"])
        lineas.append(firmante["calidad"])
        if firmante.get("es_rl"):
            lineas.append("Acepto cargo como representante legal principal.")
        elif firmante.get("es_rl_suplente"):
            lineas.append("Acepto cargo como representante legal suplente.")

        indice_nombre = 4          # las cuatro primeras son el espacio de firma
        for i, texto in enumerate(lineas):
            p = _make_paragraph(
                texto, bold=(i == indice_nombre),   # el nombre propio, resaltado
                left=True, font_name="Cambria",
                keep_next=(i < len(lineas) - 1),    # el último no arrastra
            )
            insert_point.addprevious(p)

    # ── Eliminar el párrafo residual "En calidad de poderdante" ──
    # Este párrafo del template servía como marcador para insertar las firmas
    # encima. Una vez insertadas, debe removerse para que no aparezca como
    # leyenda huérfana debajo del bloque de firma del RL.
    if poderdante_elem is not None:
        parent = poderdante_elem.getparent()
        if parent is not None:
            parent.remove(poderdante_elem)


def _build_firmantes_list(accionistas, rl_principal, rl_suplente, nombre_sas,
                          rl_principales=None, rl_suplentes=None):
    """Construye la lista de firmantes para el bloque de firmas.

    Cada firmante tiene banderas:
    - `es_rl`: representante legal principal → "Acepto cargo como RL principal."
    - `es_rl_suplente`: RL suplente → "Acepto cargo como RL suplente."
    """
    firmantes = []
    rl_nombre = rl_principal.get("nombre", "").upper()
    rls_nombre = (rl_suplente.get("nombre", "").upper()
                  if rl_suplente and rl_suplente.get("nombre") else "")

    if len(accionistas) == 1:
        acc = accionistas[0]
        nombre = acc["nombre"].upper()
        id_tipo = acc.get("id_tipo", "C.C.")
        id_num = acc.get("id_num", "")

        if nombre == rl_nombre:
            firmantes.append({
                "nombre": nombre,
                "documento": f"{id_tipo} No. {id_num}",
                "calidad": f"En calidad de constituyente único y representante legal "
                           f"principal de {nombre_sas}",
                "es_rl": True,
            })
        else:
            firmantes.append({
                "nombre": nombre,
                "documento": f"{id_tipo} No. {id_num}",
                "calidad": f"En calidad de constituyente único de {nombre_sas}",
                "es_rl": False,
                "es_rl_suplente": nombre == rls_nombre,
            })
            firmantes.append({
                "nombre": rl_nombre,
                "documento": f"C.C. No. {rl_principal.get('cc', '')}",
                "calidad": f"En calidad de representante legal principal de {nombre_sas}",
                "es_rl": True,
            })
    else:
        for acc in accionistas:
            tipo = acc.get("tipo", "natural")
            if tipo == "juridica":
                firmante_nombre = acc.get("rl_nombre", "").upper()
                firmantes.append({
                    "nombre": firmante_nombre,
                    "documento": f"C.C. No. {acc.get('rl_cc', '')}",
                    "calidad": f"En calidad de representante legal de "
                               f"{acc['nombre'].upper()}, constituyente de {nombre_sas}",
                    "es_rl": firmante_nombre == rl_nombre,
                    "es_rl_suplente": firmante_nombre == rls_nombre,
                })
            else:
                firmante_nombre = acc["nombre"].upper()
                firmantes.append({
                    "nombre": firmante_nombre,
                    "documento": f"{acc.get('id_tipo', 'C.C.')} No. {acc.get('id_num', '')}",
                    "calidad": f"En calidad de constituyente de {nombre_sas}",
                    "es_rl": firmante_nombre == rl_nombre,
                    "es_rl_suplente": firmante_nombre == rls_nombre,
                })

        # Si el RL principal no es ninguno de los accionistas, agregar bloque separado
        acc_nombres = [a["nombre"].upper() for a in accionistas]
        rl_nombres_pj = [a.get("rl_nombre", "").upper()
                         for a in accionistas if a.get("tipo") == "juridica"]
        if rl_nombre not in acc_nombres and rl_nombre not in rl_nombres_pj:
            firmantes.append({
                "nombre": rl_nombre,
                "documento": f"C.C. No. {rl_principal.get('cc', '')}",
                "calidad": f"En calidad de representante legal principal de {nombre_sas}",
                "es_rl": True,
            })

    # ── Representante Legal Suplente ──
    # Si hay suplente y NO es ya uno de los firmantes (accionista o RL principal),
    # agregar bloque de firma separado.
    if rls_nombre:
        ya_incluido = any(f["nombre"] == rls_nombre for f in firmantes)
        if not ya_incluido:
            firmantes.append({
                "nombre": rls_nombre,
                "documento": f"C.C. No. {rl_suplente.get('cc', '')}",
                "calidad": f"En calidad de representante legal suplente de {nombre_sas}",
                "es_rl": False,
                "es_rl_suplente": True,
            })

    # ── Representantes adicionales ──
    # Todo el bloque anterior razona sobre el primer principal y el primer
    # suplente, que son los que figuran ante las entidades. Los demás también
    # firman y aceptan su cargo, así que se agregan aquí.
    def _agregar(extras, calidad, clave):
        for rl in extras:
            nombre = (rl.get("nombre") or "").upper()
            if not nombre or any(f["nombre"] == nombre for f in firmantes):
                continue
            firmantes.append({
                "nombre": nombre,
                "documento": f"{_label_tipo_doc(rl.get('tipo_doc'))} No. {rl.get('cc', '')}",
                "calidad": f"En calidad de {calidad} de {nombre_sas}",
                "es_rl": clave == "principal",
                "es_rl_suplente": clave == "suplente",
            })

    _agregar((rl_principales or [])[1:], "representante legal principal", "principal")
    _agregar((rl_suplentes or [])[1:], "representante legal suplente", "suplente")

    return firmantes


def _get_para_text(p_elem):
    """Extrae el texto completo de un elemento <w:p> XML."""
    texts = [t.text or "" for t in p_elem.findall(f".//{qn('w:t')}")]
    return "".join(texts)


def _strip_text_emphasis(p_elem):
    """
    Elimina negrilla, cursiva y subrayado de un párrafo (tanto del rPr por
    defecto del párrafo como de cada run individual). Útil cuando un
    párrafo del template tiene formato enfático pero debe ser texto
    natural corrido.
    """
    enfasis_tags = ("w:b", "w:bCs", "w:i", "w:iCs", "w:u")

    # rPr por defecto del párrafo (en pPr)
    pPr = p_elem.find(qn("w:pPr"))
    if pPr is not None:
        rPr = pPr.find(qn("w:rPr"))
        if rPr is not None:
            for tag_name in enfasis_tags:
                for el in rPr.findall(qn(tag_name)):
                    rPr.remove(el)

    # rPr de cada run
    for run in p_elem.findall(qn("w:r")):
        rPr = run.find(qn("w:rPr"))
        if rPr is not None:
            for tag_name in enfasis_tags:
                for el in rPr.findall(qn(tag_name)):
                    rPr.remove(el)


def _replace_para_full_text(p_elem, new_text, strip_emphasis=False):
    """
    Reemplaza todo el texto de un párrafo preservando el formato del primer run.
    Elimina todos los runs extra para evitar espacios fantasma.

    Si `strip_emphasis=True`, también remueve negrilla, cursiva y subrayado
    del párrafo (para texto corrido que debe ser natural).
    """
    runs = p_elem.findall(qn("w:r"))
    if not runs:
        r = OxmlElement("w:r")
        rPr = OxmlElement("w:rPr")
        rFonts = OxmlElement("w:rFonts")
        rFonts.set(qn("w:ascii"), "Cambria")
        rFonts.set(qn("w:hAnsi"), "Cambria")
        rPr.append(rFonts)
        r.append(rPr)
        t = OxmlElement("w:t")
        t.text = new_text
        t.set(qn("xml:space"), "preserve")
        r.append(t)
        p_elem.append(r)
        if strip_emphasis:
            _strip_text_emphasis(p_elem)
        return

    first_run = runs[0]

    # Limpiar textos existentes del primer run y establecer nuevo
    for t_elem in list(first_run.findall(qn("w:t"))):
        first_run.remove(t_elem)
    t = OxmlElement("w:t")
    t.text = new_text
    t.set(qn("xml:space"), "preserve")
    first_run.append(t)

    # Eliminar runs extra
    for r_elem in runs[1:]:
        p_elem.remove(r_elem)

    if strip_emphasis:
        _strip_text_emphasis(p_elem)


# ════════════════════════════════════════════════════════════════
# APODERADO — LÓGICA CONDICIONAL
# ════════════════════════════════════════════════════════════════

def _handle_apoderado_section(doc, tiene_apoderado, apoderado, rl_principal, nombre_sas, genero_rl):
    """
    Maneja la sección de poder (apoderado) en las disposiciones transitorias.

    Si NO hay apoderado:
    - Elimina desde "Artículo Tercero Transitorio" ({{BLOQUE_PODER_APODERADA}})
      hasta "representan" (fin del último párrafo de poder)
    - Ajusta el párrafo de cierre eliminando "y poderdantes"
    - Elimina "En calidad de poderdante"

    Si SÍ hay apoderado:
    - Llena {{BLOQUE_PODER_APODERADA}} con los datos del Artículo Tercero Transitorio
    - Reemplaza los párrafos del Artículo Cuarto y continuaciones con versión singular
    """
    body = doc._element.body
    all_paras = list(body.findall(qn("w:p")))

    if not tiene_apoderado:
        _remove_apoderado_section(all_paras)
    else:
        _fill_apoderado_section(all_paras, apoderado, rl_principal, nombre_sas, genero_rl)


def _remove_apoderado_section(all_paras):
    """Elimina toda la sección de poder cuando no hay apoderado."""
    to_remove = []

    for p_elem in all_paras:
        text = _get_para_text(p_elem).strip()

        # Párrafo con token {{BLOQUE_PODER_APODERADA}} (no reemplazado)
        if "{{BLOQUE_PODER_APODERADA}}" in text:
            to_remove.append(p_elem)
        # Artículo Cuarto Transitorio
        elif text.startswith("El presente poder se otorga"):
            to_remove.append(p_elem)
        # Continuación: "Adicionalmente..."
        elif text.startswith("Adicionalmente, se autoriza expresamente"):
            to_remove.append(p_elem)
        # Continuación: "De la misma manera... representan"
        elif text.startswith("De la misma manera") and "representan" in text:
            to_remove.append(p_elem)
        # Pie de firma: "En calidad de poderdante"
        elif text == "En calidad de poderdante":
            to_remove.append(p_elem)
        # Ajustar párrafo de cierre: quitar "y poderdantes"
        elif "poderdantes" in text:
            for t_elem in p_elem.findall(f".//{qn('w:t')}"):
                if t_elem.text and "poderdantes" in t_elem.text:
                    t_elem.text = t_elem.text.replace(" y poderdantes", "")

    for p_elem in to_remove:
        parent = p_elem.getparent()
        if parent is not None:
            parent.remove(p_elem)


def _fill_apoderado_section(all_paras, apoderado, rl_principal, nombre_sas, genero_rl):
    """Llena la sección de poder con datos del apoderado (versión singular)."""
    ap_nombre = apoderado.get("nombre", "").upper()
    ap_id = apoderado.get("id_num", "")
    ap_ciudad = apoderado.get("domicilio_ciudad", "")
    ap_depto = apoderado.get("domicilio_departamento", "")

    # ── Artículo Tercero Transitorio: Otorgamiento de poder ──
    # Todos los suscriptores (constituyentes y RL) confieren el poder
    # Detectar género del apoderado
    ap_genero = _resolve_gender(apoderado.get("nombre", ""), "M")
    _ident = _genero(ap_genero, "identificado", "identificada")
    _domic = _genero(ap_genero, "domiciliado", "domiciliada")

    art3_text = (
        f"Los suscriptores de este documento, conferimos poder especial, "
        f"amplio y suficiente a: {ap_nombre}, {_ident} con la cédula "
        f"de ciudadanía nro. {ap_id} de {ap_ciudad}, "
        f"{_domic} en {ap_ciudad}, {ap_depto} "
        f"(en adelante el \"Apoderado\")."
    )

    # ── Artículo Cuarto Transitorio: Alcance del poder ──
    art4_text = (
        "El presente poder se otorga para que el Apoderado, realice todos los "
        "trámites relacionados con la constitución de la sociedad, para lo cual "
        "el Apoderado podrá realizar todas aquellas gestiones necesarias para "
        "llevar a feliz término dicho encargo, dentro de las que se encuentran "
        "el diligenciamiento, firma, firma electrónica, corrección, aclaración, "
        "modificación y/o adición de: (i) todos los formularios y documentos que "
        "deban presentarse para la constitución de la sociedad, incluido el "
        "presente acto de constitución; (ii) los formularios y documentos que "
        "deban presentarse para la inscripción del o de los establecimientos de "
        "comercio de la sociedad; (iii) los formularios y documentos que deban "
        "presentarse para la solicitud de los libros de comercio de la sociedad "
        "una vez ésta esté constituida (esto es, el libro de actas de la asamblea "
        "general de accionistas y el libro de registro de accionistas); (iv) los "
        "formularios y documentos que deban presentarse para la inscripción de la "
        "sociedad en el Registro Único Tributario (RUT) o para su actualización; "
        "y (v) los formularios y documentos que deban presentarse para el registro "
        "de la sociedad ante cualquier entidad nacional o municipal; hasta por un "
        "término de 12 meses."
    )

    # ── Continuación: Autorización adicional ──
    adicional_text = (
        "Adicionalmente, se autoriza expresamente al Apoderado, para que gestione "
        "todos los trámites administrativos ante la Dirección de Impuestos y Aduanas "
        "Nacionales (DIAN) y la cámara de comercio del domicilio social, así como "
        "cualquier otro trámite ante cualquier entidad pública o privada, o a las "
        "que sea necesario acudir para la correcta formación y constitución de una "
        "sociedad, incluyendo entre otras las siguientes: "
        "1. Constitución legal de la sociedad en Colombia; "
        "2. Solicitar, diligenciar, radicar, suscribir, firmar y presentar todos "
        "y cada uno de los formularios requeridos por la cámara de comercio para "
        "la inscripción del acto constitutivo de la sociedad en el registro mercantil "
        "y los necesarios para solicitar la devolución; "
        "3. Pagar derechos de registro y solicitar la devolución del dinero si es "
        "necesario ante la cámara de comercio correspondiente; "
        "4. Realizar presentación personal de cualquier documento que sea necesario "
        "para la constitución de la sociedad; "
        "5. Inscribir y registrar los libros de actas de asamblea y registro de "
        "miembros de la sociedad ante la cámara de comercio del domicilio social; "
        "6. Demás trámites ante la secretaría de hacienda, DIAN, cámara de comercio "
        "del domicilio social, municipio del domicilio social y cualquier otra gestión "
        "necesaria para la correcta constitución de la sociedad; "
        "7. Firmar y diligenciar el formulario de situación de control correspondiente "
        "y necesario al momento de la constitución de la sociedad; "
        "8. Elaborar y suscribir actas aclaratorias o adicionales del acta de "
        "constitución de la sociedad y aceptar cargos; "
        "9. Realizar presentación personal y autenticación a cualquier documento que "
        "sea necesario para la constitución de la sociedad; "
        "10. Demás diligencias propias para lograr el cumplimiento de las obligaciones "
        "formales en la etapa pre-operativa de la sociedad; y "
        "11. Solicitar la clave o usuario de la sociedad ante la cámara de comercio "
        "para efectos de renovación mercantil."
    )

    # ── Continuación: Sustitución del poder ──
    cierre_text = (
        "De la misma manera, el Apoderado podrá sustituir y reasumir el presente "
        "poder, y en general, podrán desarrollar todas las actividades requeridas "
        "para la adecuada gestión de los intereses que representan."
    )

    for p_elem in all_paras:
        text = _get_para_text(p_elem).strip()

        # Los artículos transitorios deben ser texto corrido natural —
        # sin negrilla, cursiva ni subrayado (regla del usuario).
        if "{{BLOQUE_PODER_APODERADA}}" in text:
            _replace_para_full_text(p_elem, art3_text, strip_emphasis=True)
        elif text.startswith("El presente poder se otorga"):
            _replace_para_full_text(p_elem, art4_text, strip_emphasis=True)
        elif text.startswith("Adicionalmente, se autoriza expresamente"):
            _replace_para_full_text(p_elem, adicional_text, strip_emphasis=True)
        elif text.startswith("De la misma manera") and "representan" in text:
            _replace_para_full_text(p_elem, cierre_text, strip_emphasis=True)


# ════════════════════════════════════════════════════════════════
# NOMBRAMIENTOS DE ÓRGANOS (Artículo Primero Transitorio)
# ════════════════════════════════════════════════════════════════

TIPO_DOC_LABEL = {
    "CC": "C.C.", "C.C.": "C.C.",
    "CE": "C.E.", "C.E.": "C.E.",
    "TI": "T.I.", "T.I.": "T.I.",
    "PASAPORTE": "Pasaporte", "PASSPORT": "Pasaporte",
    "NIT": "NIT",
}


def _label_tipo_doc(tipo):
    """Normaliza el tipo de documento a su etiqueta de estatutos."""
    if not tipo:
        return "C.C."
    return TIPO_DOC_LABEL.get(str(tipo).strip().upper(), str(tipo).strip())


def _persona_nombramiento(p):
    """Segmentos de 'JUAN PÉREZ, identificado con C.C. No. 71.234.567'.

    Devuelve una lista de (texto, negrilla) en lugar de una cadena, porque el
    nombre propio va resaltado y el resto de la frase no.
    """
    nombre = (p.get("nombre") or "").strip().upper()
    tipo = _label_tipo_doc(p.get("tipo_doc"))
    num = (p.get("id_num") or "").strip()
    genero = _resolve_gender(p.get("nombre", ""), p.get("genero", "M"))
    ident = _genero(genero, "identificado", "identificada")
    if tipo == "NIT":
        ident = "identificada"  # persona jurídica
    return [(nombre, True), (f", {ident} con {tipo} No. {num}", False)]


def _quitar_clausula_redundante(texto):
    """Quita la cláusula general de cierre del objeto social.

    La plantilla ya trae, justo después, su propia cláusula equivalente
    ("Asimismo la sociedad podrá llevar a cabo, en general, todas las
    operaciones..."). Si el objeto social redactado también la incluye, el
    documento queda con dos párrafos que dicen lo mismo.

    El generador ya no la produce, pero el redactor es un modelo y puede
    emitirla de todas formas: esto lo cubre.
    """
    if not texto:
        return texto
    marca = "asimismo"
    cierre = "llevar a cabo, en general"
    partes = str(texto).split("\n")
    limpias = []
    for parte in partes:
        p = parte.strip()
        pl = _strip_acentos(p)
        if pl.startswith(marca) and _strip_acentos(cierre) in pl:
            continue
        limpias.append(parte)
    resultado = "\n".join(limpias)

    # También puede venir pegada al final del mismo párrafo, sin salto previo.
    for separador in (". Asimismo, la sociedad podrá llevar a cabo",
                      ". Asimismo la sociedad podrá llevar a cabo"):
        pos = resultado.find(separador)
        if pos > 0:
            resultado = resultado[:pos + 1]
            break
    return resultado.strip()


def _normalizar_bloque(texto):
    """Limpia un bloque de texto de varias líneas antes de insertarlo.

    El objeto social lo redacta un modelo y puede llegar con líneas en blanco
    intercaladas, espacios al final de cada línea o un salto sobrante al
    cierre. Cualquiera de esos deja huecos en el documento justificado.
    """
    if not texto:
        return ""
    lineas = [" ".join(l.split()) for l in str(texto).splitlines()]
    # Se conserva una línea en blanco entre párrafos —es la que da el
    # espaciado del documento— pero nunca dos seguidas, ni al principio ni al
    # final del bloque.
    limpias = []
    for linea in lineas:
        if not linea and (not limpias or not limpias[-1]):
            continue
        limpias.append(linea)
    while limpias and not limpias[-1]:
        limpias.pop()
    return "\n".join(limpias)


# ════════════════════════════════════════════════════════════════
# REPRESENTANTES LEGALES (pueden ser varios)
# ════════════════════════════════════════════════════════════════

def _lista_representantes(data, clave_plural, clave_singular):
    """Normaliza los representantes legales a una lista.

    Acepta la forma nueva (una lista) y la antigua (uno solo), para que el
    resto del sistema y los paquetes ya generados sigan funcionando.
    """
    lista = data.get(clave_plural)
    if lista is None:
        uno = data.get(clave_singular)
        lista = [uno] if uno else []
    return [r for r in lista if r and (r.get("nombre") or "").strip()]


def _linea_representantes(representantes):
    """
    Enumera a los representantes legales en una sola frase.

    Uno:    'JUAN PÉREZ, identificado con C.C. No. 1, expedida en Medellín'
    Varios: '... ; y MARÍA GIL, identificada con C.C. No. 2, expedida en Cali'
    """
    partes = []
    for rl in representantes:
        genero = _resolve_gender(rl.get("nombre", ""), rl.get("genero", "M"))
        partes.append(
            f"{(rl.get('nombre') or '').upper()}, "
            f"{_genero(genero, 'identificado', 'identificada')} con "
            f"{_label_tipo_doc(rl.get('tipo_doc'))} No. {rl.get('cc', '')}, "
            f"expedida en {rl.get('expedicion', '')}"
        )
    if not partes:
        return ""
    if len(partes) == 1:
        return partes[0]
    return "; ".join(partes[:-1]) + "; y " + partes[-1]


def _frase_numero_representantes(n):
    """Reescribe el artículo 44 con el número real de representantes legales.

    La plantilla dice "La sociedad tendrá un representante legal, quien
    estará... será elegido... podrá ser reelegido... o removido". Con dos o
    más hay que concordar toda la frase, no solo el número.
    """
    if n <= 1:
        return None
    return (
        f"La sociedad tendrá {_numero_a_letras(n)} ({n}) representantes legales, "
        f"quienes estarán a cargo de la representación legal de la sociedad y que "
        f"serán elegidos por la asamblea general de accionistas para períodos de "
        f"un (1) año y podrán ser reelegidos indefinidamente o removidos en "
        f"cualquier tiempo."
    )


def _segmentos_a_texto(segmentos):
    """Aplana una lista de (texto, negrilla) a texto plano."""
    if isinstance(segmentos, str):
        return segmentos
    return "".join(t for t, _ in segmentos)


def _make_run(text, bold=False, font_name="Cambria", size=None):
    """Crea un <w:r> con texto y formato básico."""
    r = OxmlElement("w:r")
    rPr = OxmlElement("w:rPr")
    rFonts = OxmlElement("w:rFonts")
    rFonts.set(qn("w:ascii"), font_name)
    rFonts.set(qn("w:hAnsi"), font_name)
    rPr.append(rFonts)
    if bold:
        rPr.append(OxmlElement("w:b"))
    if size:
        sz = OxmlElement("w:sz")
        sz.set(qn("w:val"), str(size * 2))
        rPr.append(sz)
    r.append(rPr)
    t = OxmlElement("w:t")
    t.text = text
    t.set(qn("xml:space"), "preserve")
    r.append(t)
    return r


def _make_labeled_paragraph(ref_p_elem, label, content):
    """
    Crea un párrafo 'Etiqueta: contenido' clonando las propiedades de párrafo
    (alineación, interlineado, sangría) de un párrafo de referencia de la
    plantilla, de modo que los nombramientos nuevos se vean idénticos a los
    del representante legal.

    `content` puede ser una cadena o una lista de (texto, negrilla) para
    resaltar los nombres propios dentro de la frase.
    """
    p = OxmlElement("w:p")
    if ref_p_elem is not None:
        ref_pPr = ref_p_elem.find(qn("w:pPr"))
        if ref_pPr is not None:
            new_pPr = deepcopy(ref_pPr)
            # El rPr por defecto del párrafo puede traer negrilla del template
            rPr = new_pPr.find(qn("w:rPr"))
            if rPr is not None:
                for tag in ("w:b", "w:bCs", "w:i", "w:iCs", "w:u", "w:highlight"):
                    for el in rPr.findall(qn(tag)):
                        rPr.remove(el)
            p.append(new_pPr)
    p.append(_make_run(label, bold=True))
    if isinstance(content, str):
        p.append(_make_run(content, bold=False))
    else:
        for texto, negrilla in content:
            if texto:
                p.append(_make_run(texto, bold=negrilla))
    return p


def _bold_substring(p_elem, texto):
    """
    Pone en negrilla la primera aparición de `texto` dentro del párrafo,
    partiendo el run que lo contiene en hasta tres runs y conservando el
    formato original (fuente, tamaño, color) en los tres.

    Se usa para resaltar nombres propios en párrafos que ya venían armados
    por sustitución de tokens, donde no se puede elegir el formato al
    construirlos.
    """
    if not texto:
        return False
    for run in p_elem.findall(qn("w:r")):
        t_elems = run.findall(qn("w:t"))
        if len(t_elems) != 1:
            continue
        contenido = t_elems[0].text or ""
        pos = contenido.find(texto)
        if pos < 0:
            continue

        antes, medio, despues = (contenido[:pos], texto,
                                 contenido[pos + len(texto):])
        rPr = run.find(qn("w:rPr"))
        indice = list(p_elem).index(run)
        p_elem.remove(run)

        nuevos = []
        for fragmento, negrilla in ((antes, False), (medio, True), (despues, False)):
            if not fragmento:
                continue
            r = OxmlElement("w:r")
            if rPr is not None:
                nuevo_rPr = deepcopy(rPr)
                for tag in ("w:b", "w:bCs"):
                    for el in nuevo_rPr.findall(qn(tag)):
                        nuevo_rPr.remove(el)
                if negrilla:
                    nuevo_rPr.append(OxmlElement("w:b"))
                    nuevo_rPr.append(OxmlElement("w:bCs"))
                r.append(nuevo_rPr)
            elif negrilla:
                nuevo_rPr = OxmlElement("w:rPr")
                nuevo_rPr.append(OxmlElement("w:b"))
                nuevo_rPr.append(OxmlElement("w:bCs"))
                r.append(nuevo_rPr)
            t = OxmlElement("w:t")
            t.text = fragmento
            t.set(qn("xml:space"), "preserve")
            r.append(t)
            nuevos.append(r)

        for offset, r in enumerate(nuevos):
            p_elem.insert(indice + offset, r)
        return True
    return False


def _nombres_propios(accionistas, rl_principales, rl_suplentes, junta, revisor,
                     apoderado):
    """Reúne todos los nombres propios que deben ir resaltados.

    Se ordenan del más largo al más corto para que, cuando un nombre sea
    prefijo de otro, se resalte primero el completo.
    """
    nombres = []

    for acc in accionistas or []:
        nombres.append((acc.get("nombre") or "").upper())
        if acc.get("tipo") == "juridica":
            nombres.append((acc.get("rl_nombre") or "").upper())

    for rl in list(rl_principales or []) + list(rl_suplentes or []):
        if rl:
            nombres.append((rl.get("nombre") or "").upper())

    if junta:
        for grupo in ("principales", "suplentes"):
            for m in junta.get(grupo) or []:
                nombres.append((m.get("nombre") or "").upper())

    if revisor:
        nombres.append((revisor.get("nombre") or "").upper())
        nombres.append((revisor.get("contador_nombre") or "").upper())

    if apoderado:
        nombres.append((apoderado.get("nombre") or "").upper())

    return sorted({n.strip() for n in nombres if n.strip()},
                  key=len, reverse=True)


def _resaltar_nombres(doc, nombres):
    """Pone en negrilla los nombres propios donde aparezcan en el cuerpo.

    Recorre una sola vez el documento y, por cada párrafo, resalta el primer
    nombre de la lista que encuentre. Solo se aplica a párrafos del cuerpo:
    las tablas y los bloques que se construyen a mano ya salen resaltados
    desde su origen.
    """
    nombres = [n for n in dict.fromkeys(nombres) if n and len(n) > 2]
    if not nombres:
        return
    for p_elem in doc._element.body.findall(qn("w:p")):
        texto = _get_para_text(p_elem)
        if not texto.strip():
            continue
        for nombre in nombres:
            if nombre in texto:
                _bold_substring(p_elem, nombre)


# Rasgos tomados de la tabla de accionistas de la plantilla, para que el
# cuadro de la junta directiva se vea idéntico y no como una tabla ajena.
TBL_ANCHO_TOTAL = 9360      # twips
# Gris neutro para las filas de encabezado y de total. La plantilla traía un
# azul claro; el gris se lee mejor impreso y no compite con el texto.
TBL_RELLENO_ENCABEZADO = "D9D9D9"


def _tbl_borders(tag="w:tblBorders", sz="4", color="auto", inside=True):
    """Bordes de línea sencilla, con los mismos grosores de la plantilla."""
    borders = OxmlElement(tag)
    edges = ["top", "left", "bottom", "right"]
    if inside:
        edges += ["insideH", "insideV"]
    for edge in edges:
        el = OxmlElement(f"w:{edge}")
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), sz)
        el.set(qn("w:space"), "0")
        el.set(qn("w:color"), color)
        borders.append(el)
    return borders


def _pintar_fila(row, fill):
    """Aplica un relleno de fondo a todas las celdas de una fila."""
    for tc in row._tr.findall(qn("w:tc")):
        tcPr = tc.find(qn("w:tcPr"))
        if tcPr is None:
            tcPr = OxmlElement("w:tcPr")
            tc.insert(0, tcPr)
        for viejo in tcPr.findall(qn("w:shd")):
            tcPr.remove(viejo)
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"), "clear")
        shd.set(qn("w:color"), "auto")
        shd.set(qn("w:fill"), fill)
        tcPr.append(shd)


def _espaciar_despues_de_tablas(doc):
    """
    Garantiza un párrafo vacío después de cada tabla.

    Sin él la tabla queda pegada al artículo siguiente, que además es lo que
    Word necesita para no fusionar dos tablas contiguas.
    """
    body = doc._element.body
    for tbl in list(body.findall(qn("w:tbl"))):
        siguiente = tbl.getnext()
        ya_hay_blanco = (
            siguiente is not None
            and siguiente.tag == qn("w:p")
            and not _get_para_text(siguiente).strip()
        )
        if not ya_hay_blanco:
            tbl.addnext(_make_paragraph("", font_name="Cambria"))


def _tc_margenes():
    """<w:tcMar> con el relleno interno de celda de la plantilla."""
    mar = OxmlElement("w:tcMar")
    for lado, ancho in (("top", "60"), ("left", "100"),
                        ("bottom", "60"), ("right", "100")):
        el = OxmlElement(f"w:{lado}")
        el.set(qn("w:w"), ancho)
        el.set(qn("w:type"), "dxa")
        mar.append(el)
    return mar


def _make_table(rows, col_widths, header=True, font_name="Cambria", size=None):
    """
    Construye un <w:tbl> con bordes a partir de una lista de filas
    (cada fila es una lista de strings). `col_widths` en twips.

    La primera fila se marca como encabezado (negrilla + repetición en
    salto de página) cuando `header=True`.

    No se fija tamaño de fuente ni espaciado en las celdas: se hereda el del
    documento, que es lo que hace la tabla de accionistas de la plantilla.
    Fijarlos aquí produciría una tabla con métrica distinta a la otra.
    """
    tbl = OxmlElement("w:tbl")

    tblPr = OxmlElement("w:tblPr")
    tblW = OxmlElement("w:tblW")
    tblW.set(qn("w:w"), str(sum(col_widths)))
    tblW.set(qn("w:type"), "dxa")
    tblPr.append(tblW)
    tblPr.append(_tbl_borders())

    cellMar = OxmlElement("w:tblCellMar")
    for lado in ("left", "right"):
        el = OxmlElement(f"w:{lado}")
        el.set(qn("w:w"), "10")
        el.set(qn("w:type"), "dxa")
        cellMar.append(el)
    tblPr.append(cellMar)

    look = OxmlElement("w:tblLook")
    for k, v in (("w:val", "04A0"), ("w:firstRow", "1"), ("w:lastRow", "0"),
                 ("w:firstColumn", "1"), ("w:lastColumn", "0"),
                 ("w:noHBand", "0"), ("w:noVBand", "1")):
        look.set(qn(k), v)
    tblPr.append(look)
    tbl.append(tblPr)

    grid = OxmlElement("w:tblGrid")
    for w in col_widths:
        gc = OxmlElement("w:gridCol")
        gc.set(qn("w:w"), str(w))
        grid.append(gc)
    tbl.append(grid)

    for ri, row in enumerate(rows):
        tr = OxmlElement("w:tr")
        is_header = header and ri == 0
        if is_header:
            trPr = OxmlElement("w:trPr")
            trPr.append(OxmlElement("w:tblHeader"))
            tr.append(trPr)
        for ci, cell_text in enumerate(row):
            tc = OxmlElement("w:tc")
            tcPr = OxmlElement("w:tcPr")
            tcW = OxmlElement("w:tcW")
            tcW.set(qn("w:w"), str(col_widths[ci]))
            tcW.set(qn("w:type"), "dxa")
            tcPr.append(tcW)
            tcPr.append(_tbl_borders("w:tcBorders", sz="1",
                                     color="000000", inside=False))
            if is_header:
                shd = OxmlElement("w:shd")
                shd.set(qn("w:val"), "clear")
                shd.set(qn("w:color"), "auto")
                shd.set(qn("w:fill"), TBL_RELLENO_ENCABEZADO)
                tcPr.append(shd)
            tcPr.append(_tc_margenes())
            tc.append(tcPr)

            p = OxmlElement("w:p")
            # Una celda puede venir como texto plano o como lista de
            # (texto, negrilla), para resaltar el nombre propio.
            if isinstance(cell_text, str):
                p.append(_make_run(cell_text, bold=is_header,
                                   font_name=font_name, size=size))
            else:
                for texto, negrilla in cell_text:
                    if texto:
                        p.append(_make_run(texto, bold=is_header or negrilla,
                                           font_name=font_name, size=size))
            tc.append(p)
            tr.append(tc)
        tbl.append(tr)

    return tbl


def _insert_nombramientos_organos(doc, junta, revisor):
    """
    Inserta en el Artículo Primero Transitorio (bloque de nombramientos)
    los miembros de junta directiva y el revisor fiscal, justo después de
    la línea del representante legal suplente.

    `junta`   → {"principales": [...], "suplentes": [...]} o None
    `revisor` → dict con datos del revisor fiscal o None

    Cada persona es un dict con nombre, tipo_doc, id_num.
    """
    if not junta and not revisor:
        return

    body = doc._element.body
    anchor = None
    # La etiqueta va en plural cuando hay más de un suplente.
    for p_elem in body.findall(qn("w:p")):
        texto = _get_para_text(p_elem).strip()
        if (texto.startswith("Representante legal suplente:")
                or texto.startswith("Representantes legales suplentes:")):
            anchor = p_elem
            break
    if anchor is None:
        return

    # El bloque de nombramientos alterna párrafo y línea en blanco. Se clona la
    # línea en blanco que ya existe tras el representante legal suplente para
    # que los separadores nuevos tengan exactamente su mismo formato
    # (justificación, interlineado y sangría) y el ritmo vertical no se rompa.
    blank_ref = anchor.getnext()
    if blank_ref is None or blank_ref.tag != qn("w:p") \
            or _get_para_text(blank_ref).strip():
        blank_ref = None

    insert_after = anchor

    def _insert(el):
        """Inserta `el` precedido de una línea en blanco de separación."""
        nonlocal insert_after
        blank = deepcopy(blank_ref) if blank_ref is not None \
            else _make_paragraph("", font_name="Cambria")
        insert_after.addnext(blank)
        blank.addnext(el)
        insert_after = el

    # ── Junta directiva ──
    if junta:
        principales = junta.get("principales") or []
        suplentes = junta.get("suplentes") or []
        n_pri = len(principales)

        intro = (
            f"la sociedad tendrá una junta directiva integrada por "
            f"{_numero_a_letras(n_pri)} ({n_pri}) "
            f"{'miembro principal' if n_pri == 1 else 'miembros principales'}"
        )
        if suplentes:
            n_sup = len(suplentes)
            intro += (
                f" y {_numero_a_letras(n_sup)} ({n_sup}) "
                f"{'miembro suplente' if n_sup == 1 else 'miembros suplentes'} "
                f"de carácter nominal"
            )
        intro += ", designados así:"

        _insert(_make_labeled_paragraph(anchor, "Junta directiva: ", intro))

        if suplentes:
            rows = [["Miembros principales", "Miembros suplentes"]]
            for i in range(max(n_pri, len(suplentes))):
                rows.append([
                    _persona_nombramiento(principales[i]) if i < n_pri else "",
                    _persona_nombramiento(suplentes[i]) if i < len(suplentes) else "",
                ])

            col_widths = [TBL_ANCHO_TOTAL // 2] * 2
        else:
            rows = [["Miembros principales"]]
            rows.extend([[_persona_nombramiento(m)] for m in principales])
            col_widths = [TBL_ANCHO_TOTAL]

        _insert(_make_table(rows, col_widths))

    # ── Revisor fiscal ──
    # Va siempre precedido de su línea en blanco, que además separa la tabla
    # anterior del texto siguiente.
    if revisor:
        _insert(_make_labeled_paragraph(
            anchor, "Revisor fiscal: ", _texto_revisor(revisor)
        ))


def _texto_limitaciones(lim):
    """
    Redacta el parágrafo de limitaciones del representante legal.

    Es determinista: el usuario elige cuantía, naturaleza o ambas, y el
    órgano que autoriza. No se genera con IA, para que el texto no varíe
    entre documentos ni introduzca espaciados ajenos al resto del artículo.

    Devuelve segmentos (texto, negrilla) o None si no hay limitaciones.
    """
    if not lim or not lim.get("tiene_limitaciones"):
        return None

    ORGANOS = {
        "junta": "la junta directiva",
        "asamblea": "la asamblea general de accionistas",
    }
    frases = []

    if lim.get("limita_cuantia"):
        cuantia = str(lim.get("cuantia_smmlv") or "").strip()
        organo = ORGANOS.get(lim.get("organo_cuantia"), ORGANOS["asamblea"])
        if cuantia:
            frases.append(
                f"celebrar actos o contratos cuya cuantía exceda de {cuantia} "
                f"salarios mínimos mensuales legales vigentes, sin la autorización "
                f"previa de {organo}"
            )

    if lim.get("limita_naturaleza"):
        naturaleza = " ".join(str(lim.get("naturaleza") or "").split())
        organo = ORGANOS.get(lim.get("organo_naturaleza"), ORGANOS["asamblea"])
        if naturaleza:
            naturaleza = naturaleza.rstrip(".")
            frases.append(
                f"celebrar actos o contratos que versen sobre {naturaleza}, "
                f"sin la autorización previa de {organo}"
            )

    if not frases:
        return None

    if len(frases) == 1:
        cuerpo = frases[0]
    else:
        cuerpo = frases[0] + "; ni " + frases[1]

    return [
        ("Parágrafo Segundo: ", True),
        (f"El representante legal no podrá {cuerpo}. "
         f"Los actos que se celebren en contravención de esta limitación no "
         f"obligarán a la sociedad, en los términos del artículo 196 del Código "
         f"de Comercio.", False),
    ]


def _insertar_limitaciones(doc, ref_p_elem, limitaciones):
    """
    Inserta el parágrafo de limitaciones después del Parágrafo Primero del
    artículo de funciones del representante legal.

    Se clona el párrafo de referencia para heredar alineación, interlineado y
    sangría: nada de espaciados propios que rompan la homogeneidad del texto.
    """
    segmentos = _texto_limitaciones(limitaciones)
    if not segmentos:
        return False

    body = doc._element.body
    ancla = None
    for p_elem in body.findall(qn("w:p")):
        if _get_para_text(p_elem).strip().startswith("Parágrafo Primero: El representante legal"):
            ancla = p_elem
            break
    if ancla is None:
        return False

    # Línea en blanco de separación, igual a la que ya usa el documento
    blanco = ancla.getnext()
    if blanco is not None and blanco.tag == qn("w:p") and not _get_para_text(blanco).strip():
        separador = deepcopy(blanco)
    else:
        separador = _make_paragraph("", font_name="Cambria")

    parrafo = _make_labeled_paragraph(ancla, segmentos[0][0], segmentos[1:])
    ancla.addnext(parrafo)
    parrafo.addprevious(separador)
    return True


def _texto_revisor(revisor):
    """Segmentos del nombramiento del revisor fiscal (persona natural o jurídica).

    Devuelve una lista de (texto, negrilla): los nombres propios —la firma y
    el contador designado— van resaltados.
    """
    tipo = (revisor.get("tipo") or "natural").lower()
    tarjeta = (revisor.get("tarjeta_profesional") or "").strip()

    if tipo == "juridica":
        razon = (revisor.get("nombre") or "").strip().upper()
        nit = (revisor.get("id_num") or "").strip()
        c_nombre = (revisor.get("contador_nombre") or "").strip().upper()
        c_tipo = _label_tipo_doc(revisor.get("contador_tipo_doc"))
        c_num = (revisor.get("contador_id_num") or "").strip()
        c_tarjeta = (revisor.get("contador_tarjeta_profesional") or "").strip()
        c_genero = _resolve_gender(revisor.get("contador_nombre", ""), "M")
        c_ident = _genero(c_genero, "identificado", "identificada")
        c_port = _genero(c_genero, "portador", "portadora")

        return [
            (razon, True),
            (f", sociedad identificada con NIT {nit}, quien de conformidad con "
             f"el artículo 215 del Código de Comercio designa para el ejercicio "
             f"personal del cargo a ", False),
            (c_nombre, True),
            (f", {c_ident} con {c_tipo} No. {c_num}, contador público "
             f"{c_port} de la tarjeta profesional No. {c_tarjeta}.", False),
        ]

    nombre = (revisor.get("nombre") or "").strip().upper()
    tipo_doc = _label_tipo_doc(revisor.get("tipo_doc"))
    num = (revisor.get("id_num") or "").strip()
    genero = _resolve_gender(revisor.get("nombre", ""), revisor.get("genero", "M"))
    ident = _genero(genero, "identificado", "identificada")
    port = _genero(genero, "portador", "portadora")
    return [
        (nombre, True),
        (f", {ident} con {tipo_doc} No. {num}, contador público "
         f"{port} de la tarjeta profesional No. {tarjeta}.", False),
    ]


def _make_paragraph(text, bold=False, left=True, font_name="Cambria", size=None,
                    keep_next=False):
    """Crea un elemento <w:p> con formato básico.

    `keep_next` mantiene el párrafo unido al siguiente al paginar. Se usa en
    los bloques de firma para que el nombre no quede huérfano al final de una
    página y el documento no lo separe de su cédula y su calidad.
    """
    p = OxmlElement("w:p")

    # Propiedades del párrafo
    pPr = OxmlElement("w:pPr")
    if keep_next:
        pPr.append(OxmlElement("w:keepNext"))
        pPr.append(OxmlElement("w:keepLines"))
    if left:
        jc = OxmlElement("w:jc")
        jc.set(qn("w:val"), "left")
        pPr.append(jc)
    # Interlineado: sencillo sin espacio extra
    spacing = OxmlElement("w:spacing")
    spacing.set(qn("w:after"), "0")
    spacing.set(qn("w:line"), "276")  # ~1.15 líneas
    spacing.set(qn("w:lineRule"), "auto")
    pPr.append(spacing)
    p.append(pPr)

    # Run con texto
    r = OxmlElement("w:r")
    rPr = OxmlElement("w:rPr")
    rFonts = OxmlElement("w:rFonts")
    rFonts.set(qn("w:ascii"), font_name)
    rFonts.set(qn("w:hAnsi"), font_name)
    rPr.append(rFonts)
    if bold:
        b_elem = OxmlElement("w:b")
        rPr.append(b_elem)
    if size:
        sz = OxmlElement("w:sz")
        sz.set(qn("w:val"), str(size * 2))  # half-points
        rPr.append(sz)
    r.append(rPr)

    t = OxmlElement("w:t")
    t.text = text
    t.set(qn("xml:space"), "preserve")
    r.append(t)
    p.append(r)

    return p


# ════════════════════════════════════════════════════════════════
# GENERADOR PRINCIPAL
# ════════════════════════════════════════════════════════════════

def generar_estatutos(data, template_path, output_path):
    """
    Genera los estatutos de constitución S.A.S. a partir de la plantilla.

    Reemplaza los tokens {{...}} en todos los párrafos del documento.
    """
    doc = Document(template_path)

    # ── Habilitar silabeo automático para mejorar espaciado en texto justificado ──
    _enable_auto_hyphenation(doc)

    nombre_sas = data["nombre_sas"].upper()
    fecha = data["fecha"]
    if isinstance(fecha, str):
        from datetime import datetime
        fecha = datetime.strptime(fecha, "%Y-%m-%d").date()

    accionistas = data.get("accionistas", [])

    # Puede haber más de un representante legal principal y más de un
    # suplente. El resto del documento —poder, formularios— trabaja con el
    # primero de cada lista, que es el que firma y figura ante las entidades.
    rl_principales = _lista_representantes(data, "rl_principales", "rl_principal")
    rl_suplentes = _lista_representantes(data, "rl_suplentes", "rl_suplente")
    rl_principal = rl_principales[0] if rl_principales else {}
    rl_suplente = rl_suplentes[0] if rl_suplentes else None

    capital_autorizado = data.get("capital_autorizado", 1_000_000_000)
    capital_suscrito = data.get("capital_suscrito", 1_000_000)
    capital_pagado = data.get("capital_pagado", capital_suscrito)
    # Valor nominal por acción. El número de acciones de cada tramo de capital
    # se deriva dividiendo el monto entre este valor.
    valor_nominal = max(1, int(data.get("valor_nominal", 1) or 1))

    # ── Normalización automática de género ──
    # La detección por nombre prevalece sobre el dropdown del formulario,
    # para evitar errores como dejar "Yaneth" marcada como masculino por
    # olvidar cambiar el default.
    for rl in rl_principales + rl_suplentes:
        rl["genero"] = _resolve_gender(rl.get("nombre", ""), rl.get("genero", "M"))
    genero_rl = rl_principal.get("genero", "M")

    for acc in accionistas:
        if acc.get("tipo") == "juridica":
            # Detectar género del representante legal de la PJ
            acc["rl_genero"] = _resolve_gender(
                acc.get("rl_nombre", ""), acc.get("rl_genero", "M")
            )
        else:
            acc["genero"] = _resolve_gender(
                acc.get("nombre", ""), acc.get("genero", "M")
            )

    # Datos de apoderado (opcional)
    apoderado = data.get("apoderado", None)
    tiene_apoderado = bool(apoderado and apoderado.get("nombre"))

    # Líneas de nombramiento: enumeran a todos, no solo al primero
    rl_principal_linea = _linea_representantes(rl_principales)
    rl_suplente_linea = _linea_representantes(rl_suplentes) or "vacante"

    # Mapa de reemplazos
    replacements = {
        "{{NOMBRE_SOCIEDAD}}": nombre_sas,
        "{{DOMICILIO_MUNICIPIO}}": data.get("ciudad", data.get("municipio", "Medellín")),
        "{{DOMICILIO_DEPARTAMENTO}}": data.get("departamento", "Antioquia"),
        "{{FECHA_LITERAL}}": _fecha_literal(fecha),
        "{{BLOQUE_COMPARECENCIA_ACCIONISTAS}}": _bloque_comparecencia(accionistas),
        "{{OBJETO_SOCIAL_DESARROLLADO}}": _normalizar_bloque(
            _quitar_clausula_redundante(data.get("objeto_social", ""))),
        "{{CAPITAL_AUTORIZADO_MONTO_LETRAS_Y_CIFRAS}}": _monto_letras_cifras(capital_autorizado),
        "{{CAPITAL_AUTORIZADO_NUM_ACCIONES_LETRAS_Y_CIFRAS}}": _acciones_letras_cifras(capital_autorizado // valor_nominal),
        "{{CAPITAL_SUSCRITO_MONTO_LETRAS_Y_CIFRAS}}": _monto_letras_cifras(capital_suscrito),
        "{{CAPITAL_SUSCRITO_NUM_ACCIONES_LETRAS_Y_CIFRAS}}": _acciones_letras_cifras(capital_suscrito // valor_nominal),
        "{{CAPITAL_PAGADO_MONTO_LETRAS_Y_CIFRAS}}": _monto_letras_cifras(capital_pagado),
        "{{CAPITAL_PAGADO_NUM_ACCIONES_LETRAS_Y_CIFRAS}}": _acciones_letras_cifras(capital_pagado // valor_nominal),
        # Valor nominal por acción: la plantilla lo trae escrito en duro como
        # "un peso ... (COP $1)" en los artículos 4, 5 y 6. Se sustituye la
        # frase literal para que refleje el valor elegido por el usuario.
        "un peso moneda legal colombiana (COP $1)":
            _valor_nominal_frase(valor_nominal, "moneda legal colombiana"),
        "un peso colombiano (COP $1)":
            _valor_nominal_frase(valor_nominal, "colombiano"),
        # Con varios representantes se sustituye la secuencia completa de
        # tokens por la enumeración; va antes que los tokens sueltos para
        # ganarles, porque los reemplazos se aplican en orden.
        "{{RL_PRINCIPAL_NOMBRE}}, {{RL_PRINCIPAL_IDENTIFICADO}} con "
        "{{RL_PRINCIPAL_TIPO_DOC}} No. {{RL_PRINCIPAL_NUM_DOC}}, expedida en "
        "{{RL_PRINCIPAL_CIUDAD_EXPEDICION}}": rl_principal_linea,
        "{{RL_PRINCIPAL_NOMBRE}}": rl_principal.get("nombre", "").upper(),
        "{{RL_PRINCIPAL_IDENTIFICADO}}": _genero(genero_rl, "identificado", "identificada"),
        # El formulario envía "CC"/"CE"; se normaliza a la etiqueta de
        # estatutos para que la línea del representante legal no quede escrita
        # distinto a los nombramientos de junta y revisoría que van debajo.
        "{{RL_PRINCIPAL_TIPO_DOC}}": _label_tipo_doc(rl_principal.get("tipo_doc")),
        "{{RL_PRINCIPAL_NUM_DOC}}": rl_principal.get("cc", ""),
        "{{RL_PRINCIPAL_CIUDAD_EXPEDICION}}": rl_principal.get("expedicion", ""),
        "{{RL_SUPLENTE_LINEA}}": rl_suplente_linea,
        # {{BLOQUE_PODER_APODERADA}} — NO va aquí; se maneja en _handle_apoderado_section()
        "{{INICIO_TEXTO_PODER}}": "",
        "{{FIN_TEXTO_PODER}}": "",
        "{{BLOQUE_FIRMAS_FINALES}}": "",  # Se maneja con párrafos propios
    }

    # Concordancia cuando hay más de un representante legal: la etiqueta del
    # nombramiento y la frase del artículo 44 pasan a plural.
    if len(rl_principales) > 1:
        replacements["Representante legal principal:"] = "Representantes legales principales:"
    if len(rl_suplentes) > 1:
        replacements["Representante legal suplente:"] = "Representantes legales suplentes:"

    # Reemplazar en todos los párrafos del documento
    _replace_in_doc(doc, replacements)

    # El artículo de nombramiento del representante legal no tiene tokens, así
    # que su frase se cambia aparte para que diga cuántos hay.
    if len(rl_principales) > 1:
        _reemplazar_frase(
            doc,
            "La sociedad tendrá un representante legal, quien estará a cargo de la "
            "representación legal de la sociedad y que será elegido por la asamblea "
            "general de accionistas para períodos de un (1) año y podrá ser reelegido "
            "indefinidamente o removido en cualquier tiempo.",
            _frase_numero_representantes(len(rl_principales)),
        )

    # Manejar sección de apoderado (eliminar o llenar según el caso)
    _handle_apoderado_section(doc, tiene_apoderado, apoderado, rl_principal, nombre_sas, genero_rl)

    # Cada salto de línea pasa a ser un párrafo propio. Esto da el espaciado
    # entre accionistas y, sobre todo, evita que Word estire la última línea
    # de cada bloque justificado.
    _split_comparecencia_paragraph(doc)

    # Post-procesamiento estético:
    # - Título sociedad: negro (no rojo)
    # - Quitar highlight amarillo de artículos numerados
    # - Eliminar párrafos vacíos consecutivos (espaciado armónico)
    _post_process_doc(doc, nombre_sas)

    # Limitaciones del representante legal, en el artículo de sus funciones
    _insertar_limitaciones(doc, None, data.get("limitaciones_rl"))

    # Nombramientos de junta directiva y revisor fiscal en el Artículo
    # Primero Transitorio (van después del representante legal suplente)
    _insert_nombramientos_organos(
        doc, data.get("junta_directiva"), data.get("revisor_fiscal")
    )

    # ── Nombres propios en negrilla ──
    # La comparecencia, las líneas de representante legal y el poder se arman
    # por sustitución de tokens, así que el resaltado se aplica después sobre
    # el texto ya montado. Se hace antes de las firmas porque esos bloques ya
    # salen resaltados desde su construcción.
    _resaltar_nombres(doc, _nombres_propios(
        accionistas, rl_principales, rl_suplentes,
        data.get("junta_directiva"), data.get("revisor_fiscal"), apoderado,
    ))

    # Insertar firmas como párrafos individuales con espacio para firma
    _insert_firmas_paragraphs(doc, accionistas, rl_principal, rl_suplente, nombre_sas,
                              rl_principales, rl_suplentes)

    # Llenar tabla de accionistas / acciones / capital
    _fill_accionistas_table(doc, accionistas, capital_suscrito, capital_pagado,
                            valor_nominal)

    # Ninguna tabla queda pegada al artículo siguiente
    _espaciar_despues_de_tablas(doc)

    # Último paso: un solo espacio entre palabras en todo el documento. Va al
    # final para que alcance también al texto insertado por los pasos previos.
    _colapsar_espacios_dobles(doc)

    doc.save(output_path)


def _replace_in_doc(doc, replacements):
    """Reemplaza tokens en párrafos, encabezados, pies y tablas."""
    # Párrafos del cuerpo
    for para in doc.paragraphs:
        _replace_in_paragraph(para, replacements)

    # Tablas
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    _replace_in_paragraph(para, replacements)

    # Encabezados y pies de página
    for section in doc.sections:
        for header in [section.header, section.first_page_header]:
            if header:
                for para in header.paragraphs:
                    _replace_in_paragraph(para, replacements)
        for footer in [section.footer, section.first_page_footer]:
            if footer:
                for para in footer.paragraphs:
                    _replace_in_paragraph(para, replacements)


def _set_run_text_with_breaks(run, text):
    """
    Establece el texto de un run, convirtiendo \\n en elementos <w:br/>
    para que Word los renderice como saltos de línea reales.

    Word ignora los caracteres \\n dentro de <w:t>; solo <w:br/> produce
    un salto de línea visible dentro del mismo párrafo.
    """
    if "\n" not in text:
        run.text = text
        return

    r_elem = run._element
    # Eliminar elementos <w:t> existentes (python-docx crea uno al asignar .text)
    for t_elem in list(r_elem.findall(qn("w:t"))):
        r_elem.remove(t_elem)

    parts = text.split("\n")
    for i, part in enumerate(parts):
        if i > 0:
            br = OxmlElement("w:br")
            r_elem.append(br)
        t = OxmlElement("w:t")
        t.text = part
        t.set(qn("xml:space"), "preserve")
        r_elem.append(t)


def _replace_in_paragraph(para, replacements):
    """
    Reemplaza tokens en un párrafo preservando la estructura de formato.

    Estrategia:
    - Párrafos mixtos ("Subtítulo: {{TOKEN}}..."): el prefijo antes del
      primer {{ conserva su formato (negrilla), y el contenido reemplazado
      se pone en formato normal (sin negrilla).
    - Párrafos de solo token ("{{TOKEN}}"): si el texto es largo (>100 chars
      o multilínea) se quita negrilla (es texto de cuerpo). Si es corto
      (nombre propio, sociedad) conserva el formato original.

    Importante:
    - Los runs extra se eliminan del XML (no solo se vacían) para evitar
      que Word renderice espacios fantasma entre caracteres.
    - Los saltos de línea (\\n) se convierten en <w:br/> para Word.
    """
    full_text = para.text
    if "{{" not in full_text:
        return

    new_text = full_text
    for token, value in replacements.items():
        if token in new_text:
            new_text = new_text.replace(token, str(value))

    # Solo reescribir si hubo cambios
    if new_text == full_text:
        return

    if not para.runs:
        para.text = new_text
        return

    first_run = para.runs[0]
    orig_bold = first_run.bold
    first_brace = full_text.find("{{")
    p_element = para._element

    # ── Eliminar TODOS los runs extra del XML (no solo vaciar su texto) ──
    # Esto evita que Word renderice espacios fantasma por runs vacíos
    # con propiedades de formato (w:b, w:bCs, w:rFonts, etc.)
    extra_run_elements = [run._element for run in para.runs[1:]]
    for r_elem in extra_run_elements:
        p_element.remove(r_elem)

    if first_brace > 0 and orig_bold:
        # ── Párrafo mixto: "Subtítulo bold: contenido normal" ──
        # Preferir split en ": " (subtítulo: contenido) si existe antes del {{
        # para que solo el subtítulo quede en negrita.
        # Ej: "Nombre...: La sociedad se denomina {{NOMBRE}}"
        #   → "Nombre...:" en bold, "La sociedad se denomina ACME" sin bold.
        colon_pos = new_text.find(": ")
        if 0 < colon_pos < first_brace:
            split_pos = colon_pos + 2  # incluir ": " en el subtítulo bold
        else:
            split_pos = first_brace
        prefix = new_text[:split_pos]
        content = new_text[split_pos:]

        # Prefix en el primer run (mantiene negrilla del subtítulo)
        first_run.text = prefix

        # Contenido reemplazado en run nuevo sin negrilla
        content_run = para.add_run("")
        content_run.bold = False
        content_run.font.name = first_run.font.name or "Cambria"
        if first_run.font.size:
            content_run.font.size = first_run.font.size
        _set_run_text_with_breaks(content_run, content)
    else:
        # ── Párrafo de solo token o sin negrilla ──
        _set_run_text_with_breaks(first_run, new_text)

        # Si era bold y el contenido es largo (cuerpo, no nombre propio),
        # quitar negrilla para que sea texto natural
        if orig_bold and (len(new_text) > 100 or "\n" in new_text):
            first_run.bold = False
