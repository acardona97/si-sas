# -*- coding: utf-8 -*-
"""
Prueba de humo de los ajustes de documentos:
  1. Capital autorizado y valor nominal configurables
  2. Junta directiva en el Artículo Primero Transitorio
  3. Revisor fiscal persona natural y persona jurídica
  4. Carta de no declaración de situación de control
  5. Formato de empresa familiar (Ley 2495 de 2025)

Ejecutar:  python test_ajustes.py
"""
import os
import sys
from datetime import date

sys.path.insert(0, os.path.join(os.path.dirname(os.path.abspath(__file__)), "src"))

from docx import Document
from pypdf import PdfReader

from processors.estatutos import (
    generar_estatutos, _valor_nominal_frase, _texto_revisor,
    _segmentos_a_texto,
)
from processors.pdf_filler import generar_empresa_familiar, generar_carta_no_control

BASE = os.path.dirname(os.path.abspath(__file__))
PLANTILLAS = os.path.join(BASE, "plantillas")
OUT = os.path.join(BASE, "output", "_test_ajustes")
os.makedirs(OUT, exist_ok=True)


def texto_doc(path):
    doc = Document(path)
    partes = [p.text for p in doc.paragraphs]
    for t in doc.tables:
        for row in t.rows:
            partes.extend(c.text for c in row.cells)
    return "\n".join(partes)


def defectos_de_formato(path):
    """Busca en el .docx los dos defectos que estiran el texto justificado.

    1. Saltos de línea dentro de un párrafo: Word justifica la línea que
       antecede al salto, dejando huecos enormes entre las últimas palabras.
    2. Espacios dobles, que en justificado se ven como agujeros.
    """
    import re
    from docx.oxml.ns import qn
    doc = Document(path)
    saltos, dobles = [], []
    for i, p in enumerate(doc._element.body.findall(qn("w:p"))):
        texto = "".join(t.text or "" for t in p.findall(f".//{qn('w:t')}"))
        if p.find(f".//{qn('w:br')}") is not None:
            saltos.append((i, texto[:70]))
        m = re.search(r"\S(  +)\S", texto)
        if m:
            dobles.append((i, texto[max(0, m.start() - 30):m.end() + 30]))
    return saltos, dobles


def runs_en_negrilla(path):
    """Todos los fragmentos que quedaron en negrilla en el documento."""
    from docx.oxml.ns import qn
    doc = Document(path)
    negritas = set()

    def _recorrer(parrafos):
        for p in parrafos:
            for r in p.runs:
                rPr = r._element.find(qn("w:rPr"))
                if rPr is not None and rPr.find(qn("w:b")) is not None:
                    if r.text.strip():
                        negritas.add(r.text.strip())

    _recorrer(doc.paragraphs)
    for t in doc.tables:
        for row in t.rows:
            for c in row.cells:
                _recorrer(c.paragraphs)
    return negritas


# ════════════════════════════════════════════════════════════════
# 1. Valor nominal — concordancia gramatical
# ════════════════════════════════════════════════════════════════
def test_valor_nominal_frase():
    assert _valor_nominal_frase(1, "moneda legal colombiana") == \
        "un peso moneda legal colombiana (COP $1)"
    assert _valor_nominal_frase(100, "moneda legal colombiana") == \
        "cien pesos moneda legal colombiana (COP $100)"
    assert _valor_nominal_frase(1, "colombiano") == "un peso colombiano (COP $1)"
    assert _valor_nominal_frase(1000, "colombiano") == \
        "mil pesos colombianos (COP $1.000)"
    print("OK  valor nominal: concordancia singular/plural")


# ════════════════════════════════════════════════════════════════
# 2. Revisor fiscal — persona natural y jurídica
# ════════════════════════════════════════════════════════════════
def test_texto_revisor():
    segs_pn = _texto_revisor({
        "tipo": "natural", "nombre": "Ana Restrepo Gómez",
        "tipo_doc": "CC", "id_num": "43.123.456",
        "tarjeta_profesional": "98765-T",
    })
    pn = _segmentos_a_texto(segs_pn)
    assert "ANA RESTREPO GÓMEZ" in pn
    assert "identificada con C.C. No. 43.123.456" in pn
    assert "portadora de la tarjeta profesional No. 98765-T" in pn
    # Solo el nombre propio va en negrilla
    negritas_pn = [t for t, b in segs_pn if b]
    assert negritas_pn == ["ANA RESTREPO GÓMEZ"], negritas_pn

    segs_pj = _texto_revisor({
        "tipo": "juridica", "nombre": "Auditores Asociados S.A.S.",
        "id_num": "900.111.222-3",
        "contador_nombre": "Carlos Mesa Uribe", "contador_tipo_doc": "CC",
        "contador_id_num": "71.999.888",
        "contador_tarjeta_profesional": "12345-T",
    })
    pj = _segmentos_a_texto(segs_pj)
    assert "NIT 900.111.222-3" in pj
    assert "CARLOS MESA URIBE" in pj.upper()
    assert "identificado con C.C. No. 71.999.888" in pj
    assert "portador de la tarjeta profesional No. 12345-T" in pj
    assert "artículo 215 del Código de Comercio" in pj
    # La firma y el contador designado van resaltados; el resto no
    negritas_pj = [t for t, b in segs_pj if b]
    assert negritas_pj == ["AUDITORES ASOCIADOS S.A.S.", "CARLOS MESA URIBE"], negritas_pj
    print("OK  revisor fiscal: persona natural y jurídica")


# ════════════════════════════════════════════════════════════════
# 3. Estatutos completos — nominal $100 + junta + revisor PJ
# ════════════════════════════════════════════════════════════════
def test_estatutos():
    data = {
        "nombre_sas": "PRUEBA TEST S.A.S.",
        "fecha": date(2026, 7, 28),
        "municipio": "Medellín", "ciudad": "Medellín",
        "departamento": "Antioquia",
        "accionistas": [
            {"tipo": "natural", "nombre": "Juan Pablo García", "id_tipo": "C.C.",
             "id_num": "71111111", "expedicion": "Medellín", "domicilio": "Medellín",
             "porcentaje": 60, "genero": "M"},
            {"tipo": "natural", "nombre": "María Camila Torres", "id_tipo": "C.C.",
             "id_num": "43222222", "expedicion": "Envigado", "domicilio": "Envigado",
             "porcentaje": 40, "genero": "F"},
        ],
        # Redactado en varios párrafos, con líneas en blanco, espacios al
        # final y un salto sobrante: justo lo que devuelve el modelo.
        "objeto_social": (
            "La sociedad tendrá como objeto social principal las actividades de "
            "desarrollo de sistemas informáticos.   \n"
            "\n"
            "Adicionalmente, la sociedad podrá celebrar toda clase de actos y "
            "contratos civiles y comerciales; importar y exportar bienes y "
            "servicios; participar como socia o accionista en otras sociedades; "
            "así como obtener créditos y financiamientos de cualquier naturaleza "
            "con entidades financieras para fortalecer su operación y expansión "
            "empresarial.\n"
            "\n"
            # Cláusula que la plantilla ya trae: debe eliminarse para no dejar
            # dos párrafos diciendo lo mismo.
            "Asimismo, la sociedad podrá llevar a cabo, en general, todas las "
            "actividades lícitas de comercio que guarden relación directa o "
            "indirecta con el objeto social descrito, incluyendo aquellas "
            "actividades accesorias, complementarias o conexas que sean "
            "necesarias o convenientes para el cumplimiento de su objeto social "
            "principal.\n"
        ),
        # 500.000.000 / 100 = 5.000.000 acciones
        "capital_autorizado": 500_000_000,
        "capital_suscrito": 2_000_000,
        "capital_pagado": 1_000_000,
        "valor_nominal": 100,
        "rl_principal": {"nombre": "Juan Pablo García", "cc": "71111111",
                         "tipo_doc": "C.C.", "expedicion": "Medellín", "genero": "M"},
        "rl_suplente": None,
        "tiene_junta": True, "tiene_revisor": True,
        "junta_directiva": {
            "principales": [
                {"nombre": "Juan Pablo García", "tipo_doc": "CC", "id_num": "71111111"},
                {"nombre": "María Camila Torres", "tipo_doc": "CC", "id_num": "43222222"},
                {"nombre": "Peter Schmidt", "tipo_doc": "Pasaporte", "id_num": "X8899221"},
            ],
            "suplentes": [
                {"nombre": "Laura Gil Peña", "tipo_doc": "CC", "id_num": "43555444"},
            ],
        },
        "revisor_fiscal": {
            "tipo": "juridica", "nombre": "Auditores Asociados S.A.S.",
            "id_num": "900.111.222-3",
            "contador_nombre": "Carlos Mesa Uribe", "contador_tipo_doc": "CC",
            "contador_id_num": "71.999.888",
            "contador_tarjeta_profesional": "12345-T",
        },
        "apoderado": None,
    }
    out = os.path.join(OUT, "estatutos.docx")
    generar_estatutos(data, os.path.join(PLANTILLAS, "estatutos_template.docx"), out)
    txt = texto_doc(out)

    assert "{{" not in txt, "quedaron tokens sin reemplazar en los estatutos"

    # Artículo 4: capital autorizado y acciones derivadas del nominal
    assert "QUINIENTOS MILLONES DE PESOS (COP $500.000.000)" in txt
    assert "cinco millones (5.000.000) acciones" in txt
    assert "cien pesos moneda legal colombiana (COP $100)" in txt
    assert "un peso moneda legal colombiana (COP $1)" not in txt

    # Artículo 5 y 6: suscrito 2.000.000/100 = 20.000; pagado 1.000.000/100 = 10.000
    assert "veinte mil (20.000) acciones" in txt
    assert "diez mil (10.000) acciones" in txt
    assert "cien pesos colombianos (COP $100)" in txt

    # Artículo Primero Transitorio: junta directiva
    assert "Junta directiva:" in txt
    assert "tres (3) miembros principales" in txt
    assert "un (1) miembro suplente" in txt
    assert "Miembros principales" in txt and "Miembros suplentes" in txt
    assert "PETER SCHMIDT, identificado con Pasaporte No. X8899221" in txt
    assert "LAURA GIL PEÑA, identificada con C.C. No. 43555444" in txt

    # Artículo Primero Transitorio: revisor fiscal
    assert "Revisor fiscal:" in txt
    assert "AUDITORES ASOCIADOS S.A.S., sociedad identificada con NIT 900.111.222-3" in txt

    # La tabla de accionistas conserva sus totales pese a la tabla nueva
    doc = Document(out)
    tabla_acc = next(t for t in doc.tables
                     if t.rows and "accionista" in t.rows[0].cells[0].text.lower())
    total = tabla_acc.rows[-1]
    assert total.cells[0].text.strip() == "TOTAL"
    assert total.cells[3].text.strip() == "20.000", total.cells[3].text
    assert total.cells[4].text.strip() == "$2.000.000"

    # ── Todos los nombres propios en negrilla ──
    negritas = runs_en_negrilla(out)
    esperados = [
        "JUAN PABLO GARCÍA",        # accionista, RL, miembro de junta y firmante
        "MARÍA CAMILA TORRES",      # accionista y miembro de junta
        "PETER SCHMIDT",            # miembro de junta
        "LAURA GIL PEÑA",           # suplente de junta
        "AUDITORES ASOCIADOS S.A.S.",   # revisor fiscal persona jurídica
        "CARLOS MESA URIBE",        # contador designado
    ]
    for nombre in esperados:
        assert nombre in negritas, (
            f"{nombre!r} debía quedar en negrilla. En negrilla hay: "
            f"{sorted(negritas)[:25]}"
        )
    # El texto corriente que rodea al nombre NO debe heredar la negrilla
    assert not any(n.startswith("identificado con") or n.startswith(", identificad")
                   for n in negritas), sorted(negritas)
    print(f"OK  estatutos con nominal $100 + junta + revisor PJ -> {out}")
    print(f"OK  nombres propios en negrilla ({len(esperados)} verificados)")

    # ── Nada que estire el texto justificado ──
    saltos, dobles = defectos_de_formato(out)
    assert not saltos, (
        "quedaron saltos de línea dentro de un párrafo; Word estira la línea "
        f"anterior al salto: {saltos}"
    )
    assert not dobles, f"quedaron espacios dobles: {dobles}"

    # El objeto social multipárrafo se partió en párrafos de verdad
    parrafos = [p.text.strip() for p in Document(out).paragraphs if p.text.strip()]
    assert any(p.startswith("Adicionalmente, la sociedad podrá celebrar") for p in parrafos), \
        "el segundo párrafo del objeto social debía quedar como párrafo propio"
    assert any(p.endswith("expansión empresarial.") for p in parrafos)

    # Los tramos de continuación no pueden llevar numeración automática: si la
    # llevaran, se comerían el número del artículo siguiente.
    from docx.oxml.ns import qn as _qn
    doc_v = Document(out)
    for p in doc_v.paragraphs:
        if p.text.strip().startswith("Adicionalmente, la sociedad podrá celebrar"):
            pPr = p._element.find(_qn("w:pPr"))
            assert pPr is None or pPr.find(_qn("w:numPr")) is None, \
                "el párrafo de continuación no debe numerarse como artículo"
            break
    # Y el articulado sigue corrido: el capital autorizado conserva su artículo
    assert any(p.startswith("Capital autorizado:") for p in parrafos), \
        "se perdió el artículo de capital autorizado"
    print("OK  sin saltos internos ni espacios dobles (formato simétrico)")
    print("OK  la numeración de artículos no se altera al partir párrafos")

    # ── El objeto social no repite la cláusula de cierre de la plantilla ──
    cierres = [p for p in parrafos
               if p.lower().startswith("asimismo")
               and "llevar a cabo, en general" in p.lower()]
    assert len(cierres) == 1, (
        f"la cláusula general de cierre debe aparecer una sola vez (la de la "
        f"plantilla); se encontraron {len(cierres)}"
    )
    assert "todas las operaciones" in cierres[0], \
        "debe quedar la cláusula de la plantilla, no la del objeto social redactado"

    # ── Los cuadros: encabezado gris y separados del texto siguiente ──
    from docx.oxml.ns import qn as _q
    doc_t = Document(out)

    def _relleno(fila):
        rellenos = set()
        for tc in fila._tr.findall(_q("w:tc")):
            tcPr = tc.find(_q("w:tcPr"))
            shd = tcPr.find(_q("w:shd")) if tcPr is not None else None
            rellenos.add(shd.get(_q("w:fill")) if shd is not None else None)
        return rellenos

    tablas = [t for t in doc_t.tables if t.rows and len(t.rows) > 1]
    assert len(tablas) >= 2, "deben existir el cuadro de junta y el de accionistas"
    for t in tablas:
        assert _relleno(t.rows[0]) == {"D9D9D9"}, (
            f"el encabezado debe ir en gris, no en azul: {_relleno(t.rows[0])}"
        )
    # La fila TOTAL del cuadro de accionistas también
    tabla_acc2 = next(t for t in doc_t.tables
                      if t.rows and "accionista" in t.rows[0].cells[0].text.lower())
    assert _relleno(tabla_acc2.rows[-1]) == {"D9D9D9"}, "la fila TOTAL también va en gris"

    # Después de cada tabla debe venir un párrafo vacío
    cuerpo = list(doc_t._element.body)
    for i, el in enumerate(cuerpo):
        if el.tag != _q("w:tbl"):
            continue
        siguiente = cuerpo[i + 1] if i + 1 < len(cuerpo) else None
        assert siguiente is not None and siguiente.tag == _q("w:p"), \
            "toda tabla debe ir seguida de un párrafo"
        texto_sig = "".join(t.text or "" for t in siguiente.findall(f".//{_q('w:t')}"))
        assert not texto_sig.strip(), \
            f"la tabla no puede quedar pegada al texto siguiente: {texto_sig[:60]!r}"
    print("OK  cuadros con encabezado gris y separados del artículo siguiente")


def test_estatutos_defaults():
    """Sin junta ni revisor y con nominal $1, el documento no cambia de forma."""
    data = {
        "nombre_sas": "SIMPLE S.A.S.",
        "fecha": date(2026, 7, 28),
        "municipio": "Medellín", "ciudad": "Medellín", "departamento": "Antioquia",
        "accionistas": [
            {"tipo": "natural", "nombre": "Ana Restrepo", "id_tipo": "C.C.",
             "id_num": "43000111", "expedicion": "Medellín", "domicilio": "Medellín",
             "porcentaje": 100, "genero": "F"},
        ],
        "objeto_social": "Actividades de prueba.",
        "capital_autorizado": 1_000_000_000,
        "capital_suscrito": 1_000_000,
        "capital_pagado": 1_000_000,
        "valor_nominal": 1,
        "rl_principal": {"nombre": "Ana Restrepo", "cc": "43000111",
                         "tipo_doc": "C.C.", "expedicion": "Medellín", "genero": "F"},
        "rl_suplente": None,
        "tiene_junta": False, "tiene_revisor": False,
        "junta_directiva": None, "revisor_fiscal": None, "apoderado": None,
    }
    out = os.path.join(OUT, "estatutos_simple.docx")
    generar_estatutos(data, os.path.join(PLANTILLAS, "estatutos_template.docx"), out)
    txt = texto_doc(out)

    assert "{{" not in txt
    assert "un peso moneda legal colombiana (COP $1)" in txt
    assert "mil millones de pesos (COP $1.000.000.000)".upper() in txt.upper()
    assert "Junta directiva:" not in txt
    assert "Revisor fiscal:" not in txt
    print(f"OK  estatutos por defecto sin junta ni revisor -> {out}")


# ════════════════════════════════════════════════════════════════
# 4. Carta de no declaración de situación de control
# ════════════════════════════════════════════════════════════════
def test_carta_no_control():
    for nombre, unico in (("carta_unico.pdf", True), ("carta_mayoritario.pdf", False)):
        out = os.path.join(OUT, nombre)
        generar_carta_no_control({
            "nombre_sas": "PRUEBA TEST S.A.S.",
            "fecha": date(2026, 7, 28),
            "ciudad": "Medellín",
            "camara_ciudad": "Medellín",
            "nombre_rl": "Juan Pablo García",
            "tipo_doc_rl": "C.C.", "cc_rl": "71111111",
            "controlante_nombre": "JUAN PABLO GARCÍA",
            "controlante_porcentaje": 100 if unico else 60,
            "es_unico_accionista": unico,
        }, out)
        # El texto va justificado por líneas: se normalizan los espacios para
        # que las aserciones no dependan de dónde cayó el salto de línea ni de
        # en qué página quedó el bloque de firma.
        lector = PdfReader(out)
        txt = " ".join(" ".join(p.extract_text() or "" for p in lector.pages).split())
        assert "CÁMARA DE COMERCIO DE MEDELLÍN" in txt
        assert "PRUEBA TEST S.A.S." in txt
        assert "artículo 261 del Código de Comercio" in txt
        assert "JUAN PABLO GARCÍA" in txt
        assert "Representante legal principal" in txt
        if unico:
            assert "único accionista" in txt
        else:
            assert "accionista mayoritario" in txt
            assert "60%" in txt

        # El pie no puede encimarse sobre el bloque de firma: se comprueba que
        # ningún texto del cuerpo invada la banda inferior donde va el pie.
        posiciones = []

        def _visitor(text, cm, tm, font_dict, font_size):
            if text and text.strip():
                posiciones.append((round(tm[5], 1), round(font_size, 1), text.strip()))

        for pagina in PdfReader(out).pages:
            posiciones.clear()
            pagina.extract_text(visitor_text=_visitor)
            cuerpo = [p for p in posiciones if p[1] > 9]      # el pie va en 8 pt
            pie = [p for p in posiciones if p[1] <= 9]
            if cuerpo and pie:
                y_pie = max(p[0] for p in pie)
                y_min_cuerpo = min(p[0] for p in cuerpo)
                assert y_min_cuerpo > y_pie + 8, (
                    f"el cuerpo baja hasta y={y_min_cuerpo} y el pie está en "
                    f"y={y_pie}: se encimarían"
                )
            # Nada debe quedar pegado al borde inferior de la página
            if cuerpo:
                assert min(p[0] for p in cuerpo) > 40, "hay texto al borde de la hoja"

        # ── El cuerpo va justificado a ambos márgenes ──
        # Se agrupan los fragmentos por línea y se mide dónde termina cada una:
        # en texto justificado casi todas deben llegar al margen derecho.
        from reportlab.pdfbase.pdfmetrics import stringWidth
        piezas = []

        def _v(text, cm, tm, font_dict, font_size):
            if text and text.strip():
                piezas.append((round(tm[5], 1), tm[4], text, round(font_size, 1)))

        PdfReader(out).pages[0].extract_text(visitor_text=_v)
        MARGEN_DER = 612 - 72          # carta menos margen derecho
        lineas = {}
        for y, x, txt, sz in piezas:
            if sz <= 9:                # el pie no cuenta
                continue
            # pypdf agrega un salto a cada fragmento: se descarta para medir.
            lineas.setdefault(y, []).append(x + stringWidth(txt.strip(), "Times-Roman", sz))

        # Una línea justificada se dibuja palabra por palabra, así que llega
        # con varios fragmentos; las demás (últimas de párrafo, encabezado y
        # firma) vienen en uno solo y van alineadas a la izquierda.
        justificadas = [max(v) for v in lineas.values() if len(v) > 1]
        sueltas = [max(v) for v in lineas.values() if len(v) == 1]

        assert len(justificadas) >= 10, (
            f"casi ninguna línea salió justificada ({len(justificadas)}): "
            "el cuerpo de la carta debe ir justificado a ambos márgenes"
        )
        desalineadas = [f for f in justificadas if abs(f - MARGEN_DER) > 3]
        assert not desalineadas, (
            f"hay líneas justificadas que no llegan al margen derecho "
            f"({MARGEN_DER}): {[round(f, 1) for f in desalineadas]}"
        )
        # Ninguna línea puede desbordarse del margen
        assert max(justificadas + sueltas) < MARGEN_DER + 3, \
            f"hay texto fuera del margen: {max(justificadas + sueltas):.1f}"

        print(f"OK  carta no situación de control ({'único' if unico else 'mayoritario'}) -> {out}")
        print(f"    justificada: {len(justificadas)} líneas al margen exacto, "
              f"{len(sueltas)} de cierre de párrafo")


# ════════════════════════════════════════════════════════════════
# 5. Formato de empresa familiar (Ley 2495 de 2025)
# ════════════════════════════════════════════════════════════════
def test_empresa_familiar():
    out = os.path.join(OUT, "empresa_familiar.pdf")
    generar_empresa_familiar({
        "nombre_sas": "PRUEBA TEST S.A.S.",
        "fecha": date(2026, 7, 28),
        "ciudad": "Medellín", "camara_ciudad": "Medellín",
        "nombre_rl": "Juan Pablo García",
        "tipo_doc_rl": "C.C.", "cc_rl": "71111111",
        "nucleo_familiar": [
            {"tipo_doc": "C.C.", "id_num": "71111111",
             "nombre": "Juan Pablo García", "acciones": "1.200.000",
             "parentesco": "Padre"},
            {"tipo_doc": "C.C.", "id_num": "43222222",
             "nombre": "María Camila Torres Villegas", "acciones": "800.000",
             "parentesco": "Hija"},
        ],
    }, os.path.join(PLANTILLAS, "empresa_familiar_form.pdf"), out)

    txt = PdfReader(out).pages[0].extract_text()
    assert "LEY 2495 DE 2025" in txt, "se perdió el texto original del formato"
    assert "MEDELLÍN" in txt.upper()
    assert "PRUEBA TEST S.A.S." in txt
    assert "JUAN PABLO GARCÍA" in txt.upper()
    assert "71111111" in txt
    assert "Padre" in txt and "Hija" in txt
    assert "1.200.000" in txt and "800.000" in txt
    print(f"OK  formato empresa familiar -> {out}")


def test_varios_representantes_y_limitaciones():
    """Dos principales, dos suplentes y limitaciones de cuantía y naturaleza."""
    base = {
        "nombre_sas": "MULTI RL S.A.S.",
        "fecha": date(2026, 7, 29),
        "municipio": "Medellín", "ciudad": "Medellín", "departamento": "Antioquia",
        "accionistas": [
            {"tipo": "natural", "nombre": "Ana Restrepo", "id_tipo": "C.C.",
             "id_num": "43000111", "expedicion": "Medellín", "domicilio": "Medellín",
             "porcentaje": 100, "genero": "F"},
        ],
        "objeto_social": "Actividades de prueba.",
        "capital_autorizado": 1_000_000_000,
        "capital_suscrito": 1_000_000, "capital_pagado": 1_000_000, "valor_nominal": 1,
        "rl_principales": [
            {"nombre": "Ana Restrepo", "cc": "43000111", "tipo_doc": "CC",
             "expedicion": "Medellín", "genero": "F"},
            {"nombre": "Pedro Gómez", "cc": "71222333", "tipo_doc": "CC",
             "expedicion": "Envigado", "genero": "M"},
        ],
        "rl_suplentes": [
            {"nombre": "Luisa Mora", "cc": "43555666", "tipo_doc": "CC",
             "expedicion": "Medellín", "genero": "F"},
            {"nombre": "Carlos Ruiz", "cc": "71888999", "tipo_doc": "CC",
             "expedicion": "Itagüí", "genero": "M"},
        ],
        "limitaciones_rl": {
            "tiene_limitaciones": True,
            "limita_cuantia": True, "cuantia_smmlv": "500", "organo_cuantia": "asamblea",
            "limita_naturaleza": True,
            "naturaleza": "la enajenación o gravamen de bienes inmuebles",
            "organo_naturaleza": "junta",
        },
        "tiene_junta": False, "tiene_revisor": False,
        "junta_directiva": None, "revisor_fiscal": None, "apoderado": None,
    }
    out = os.path.join(OUT, "estatutos_multi_rl.docx")
    generar_estatutos(base, os.path.join(PLANTILLAS, "estatutos_template.docx"), out)
    txt = texto_doc(out)
    assert "{{" not in txt

    # ── Artículo 44: el número real de representantes legales ──
    assert "La sociedad tendrá dos (2) representantes legales" in txt, \
        "el artículo 44 debe decir cuántos representantes legales hay"
    assert "La sociedad tendrá un representante legal," not in txt
    # Y concuerda toda la frase, no solo el número
    assert "quienes estarán a cargo" in txt
    assert "serán elegidos por la asamblea" in txt
    assert "podrán ser reelegidos indefinidamente o removidos" in txt
    # La frase de los suplentes de la plantilla se conserva
    assert "podrá tener uno o varios suplentes" in txt

    # ── Nombramientos: los cuatro, con etiqueta en plural ──
    assert "Representantes legales principales:" in txt
    assert "Representantes legales suplentes:" in txt
    for nombre in ["ANA RESTREPO", "PEDRO GÓMEZ", "LUISA MORA", "CARLOS RUIZ"]:
        assert nombre in txt, f"falta {nombre} en los nombramientos"
    assert "; y PEDRO GÓMEZ" in txt, "los principales deben ir enumerados"

    # ── Todos firman ──
    firmas = [p.text.strip() for p in Document(out).paragraphs]
    for nombre in ["PEDRO GÓMEZ", "CARLOS RUIZ"]:
        assert nombre in firmas, f"{nombre} debía tener bloque de firma propio"
    assert "Acepto cargo como representante legal principal." in firmas
    assert "Acepto cargo como representante legal suplente." in firmas

    # ── Limitaciones en el artículo de funciones ──
    assert "Parágrafo Segundo:" in txt
    assert "500 salarios mínimos mensuales legales vigentes" in txt
    assert "sin la autorización previa de la asamblea general de accionistas" in txt
    assert "la enajenación o gravamen de bienes inmuebles" in txt
    assert "sin la autorización previa de la junta directiva" in txt
    assert "artículo 196 del Código de Comercio" in txt

    # ── Y sin romper la forma ──
    saltos, dobles = defectos_de_formato(out)
    assert not saltos, saltos
    assert not dobles, dobles

    # El parágrafo nuevo hereda el formato del Parágrafo Primero
    from docx.oxml.ns import qn as _q
    d = Document(out)
    p1 = next(p for p in d.paragraphs
              if p.text.strip().startswith("Parágrafo Primero: El representante legal"))
    p2 = next(p for p in d.paragraphs if p.text.strip().startswith("Parágrafo Segundo:"))

    def _props(p):
        pPr = p._element.find(_q("w:pPr"))
        if pPr is None:
            return (None, None, None)
        jc = pPr.find(_q("w:jc"))
        sp = pPr.find(_q("w:spacing"))
        ind = pPr.find(_q("w:ind"))
        return (
            jc.get(_q("w:val")) if jc is not None else None,
            (sp.get(_q("w:after")), sp.get(_q("w:line"))) if sp is not None else None,
            (ind.get(_q("w:left")), ind.get(_q("w:firstLine"))) if ind is not None else None,
        )

    assert _props(p2) == _props(p1), (
        f"el parágrafo de limitaciones debe tener el mismo formato que el "
        f"Parágrafo Primero: {_props(p2)} vs {_props(p1)}"
    )
    print(f"OK  varios representantes legales y limitaciones -> {out}")
    print("OK  el parágrafo de limitaciones hereda el formato del anterior")


def test_un_solo_representante_no_cambia():
    """Con un principal y un suplente el texto queda como siempre."""
    txt = texto_doc(os.path.join(OUT, "estatutos_simple.docx"))
    assert "La sociedad tendrá un representante legal," in txt
    assert "Representante legal principal:" in txt
    assert "Representantes legales principales:" not in txt
    assert "Parágrafo Segundo:" not in txt, \
        "sin limitaciones no debe aparecer el parágrafo"
    print("OK  con un solo representante y sin limitaciones nada cambia")


if __name__ == "__main__":
    test_valor_nominal_frase()
    test_texto_revisor()
    test_estatutos()
    test_estatutos_defaults()
    test_varios_representantes_y_limitaciones()
    test_un_solo_representante_no_cambia()
    test_carta_no_control()
    test_empresa_familiar()
    print("\nTodas las pruebas pasaron.")
