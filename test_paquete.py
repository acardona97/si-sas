# -*- coding: utf-8 -*-
"""
Prueba de integración del paquete completo: arma el mismo JSON que envía el
cuestionario y verifica los documentos que salen del ZIP de /api/generate.

Escenarios:
  A. Dos accionistas, controlante al 60% que NO declara control, junta
     directiva con suplentes, revisor fiscal persona jurídica, empresa
     familiar, capital autorizado y valor nominal personalizados.
  B. Accionista único que SÍ declara control, sin junta ni revisor,
     valores por defecto.

Ejecutar:  python test_paquete.py
"""
import io
import os
import sys
import zipfile

sys.path.insert(0, os.path.join(os.path.dirname(os.path.abspath(__file__)), "src"))

from pypdf import PdfReader
from docx import Document

import app as flask_app
import auth

BASE = os.path.dirname(os.path.abspath(__file__))
OUT = os.path.join(BASE, "output", "_test_paquete")
os.makedirs(OUT, exist_ok=True)

TEST_EMAIL = "prueba-paquete@test.co"


def _cliente():
    """Cliente Flask autenticado con generaciones ilimitadas."""
    auth.init_db()
    user = auth.get_user_by_email(TEST_EMAIL)
    if not user:
        user = auth.create_user(TEST_EMAIL, "clave-de-prueba-123", "Prueba Paquete")
    auth.set_generaciones(user["id"], None)  # ilimitado

    flask_app.app.config["TESTING"] = True
    client = flask_app.app.test_client()
    with client.session_transaction() as sess:
        sess["user_id"] = user["id"]
        sess["user_email"] = user["email"]
        sess["plan"] = user["plan"]
    return client


def _generar(client, payload, carpeta):
    resp = client.post("/api/generate", json=payload)
    assert resp.status_code == 200, f"HTTP {resp.status_code}: {resp.get_data(as_text=True)[:600]}"

    destino = os.path.join(OUT, carpeta)
    os.makedirs(destino, exist_ok=True)
    archivos = {}
    with zipfile.ZipFile(io.BytesIO(resp.data)) as zf:
        for nombre in zf.namelist():
            ruta = os.path.join(destino, nombre)
            with open(ruta, "wb") as f:
                f.write(zf.read(nombre))
            archivos[nombre] = ruta
    return archivos


PAYLOAD_DIR = os.path.join(BASE, "output", "_test_cuestionario")


def _payload(nombre, respaldo):
    """Usa el payload que volcó test_cuestionario.mjs si existe.

    Así se verifica el JSON real que arma el cuestionario y no una copia a
    mano que puede quedar desactualizada. Si no se ha corrido el test de
    Node, se cae al payload de respaldo.
    """
    import json
    ruta = os.path.join(PAYLOAD_DIR, nombre)
    if os.path.exists(ruta):
        with open(ruta, encoding="utf-8") as f:
            print(f"  (usando el payload real del cuestionario: {nombre})")
            return json.load(f)
    print(f"  (aviso: no existe {nombre}; corra 'node test_cuestionario.mjs' "
          f"para verificar el payload real)")
    return respaldo


def _pdf_txt(path):
    """Texto de un PDF con espacios normalizados."""
    partes = [p.extract_text() or "" for p in PdfReader(path).pages]
    return " ".join(" ".join(partes).split())


def _docx_txt(path):
    doc = Document(path)
    partes = [p.text for p in doc.paragraphs]
    for t in doc.tables:
        for row in t.rows:
            partes.extend(c.text for c in row.cells)
    return "\n".join(partes)


def _buscar(archivos, fragmento):
    for nombre, ruta in archivos.items():
        if fragmento in nombre:
            return ruta
    return None


# ════════════════════════════════════════════════════════════════
# ESCENARIO A
# ════════════════════════════════════════════════════════════════
PAYLOAD_A = {
    "nombre_sas": "FAMILIA ANDINA S.A.S.",
    "municipio": "Medellín", "departamento": "Antioquia",
    "direccion": "Carrera 43A #1-50, Oficina 805", "barrio": "El Poblado",
    "email": "contacto@familiaandina.co",
    "telefono1": "3001234567", "telefono2": "", "telefono3": "",
    "zona": "urbana", "tipo_local": "oficina", "tenencia": "arriendo",
    "accionistas": [
        {"tipo": "natural", "nombre": "Juan Pablo García", "tipo_doc": "CC",
         "id_tipo": "C.C.", "id_num": "71111111", "expedicion": "Medellín",
         "domicilio": "Medellín", "nacimiento": "", "genero": "M",
         "porcentaje": 60, "capital_pagado": ""},
        {"tipo": "natural", "nombre": "María Camila Torres", "tipo_doc": "CC",
         "id_tipo": "C.C.", "id_num": "43222222", "expedicion": "Envigado",
         "domicilio": "Envigado", "nacimiento": "", "genero": "F",
         "porcentaje": 40, "capital_pagado": ""},
    ],
    "rl_principal": {"nombre": "Juan Pablo García", "cc": "71111111",
                     "tipo_doc": "C.C.", "expedicion": "Medellín", "genero": "M"},
    "rl_suplente": None,
    "ciiu_code": "6201", "ciiu_description": "Actividades de desarrollo de sistemas informáticos",
    "ciiu_code_sec": "", "ciiu_description_sec": "",
    "objeto_social": "Desarrollo de software a la medida.",
    # 500.000.000 / 100 = 5.000.000 acciones autorizadas
    "capital_autorizado": "500.000.000",
    "valor_nominal": "100",
    "capital_suscrito": "2.000.000",
    "capital_pagado": "",
    "regimen": "ordinario", "ingresos_mensuales": "100.000",
    "tiene_junta": True, "tiene_revisor": True,
    "junta_directiva": {
        "principales": [
            {"nombre": "Juan Pablo García", "tipo_doc": "CC", "id_num": "71111111"},
            {"nombre": "María Camila Torres", "tipo_doc": "CC", "id_num": "43222222"},
            {"nombre": "Peter Schmidt", "tipo_doc": "Pasaporte", "id_num": "X8899221"},
        ],
        "suplentes": [
            {"nombre": "Laura Gil Peña", "tipo_doc": "CC", "id_num": "43555444"},
        ],
    },
    "revisor_fiscal": {
        "tipo": "juridica", "nombre": "Auditores Asociados S.A.S.",
        "id_num": "900.111.222-3",
        "contador_nombre": "Carlos Mesa Uribe", "contador_tipo_doc": "CC",
        "contador_id_num": "71999888", "contador_tarjeta_profesional": "12345-T",
    },
    "es_emprendimiento_social": False, "grupo_etnico": "",
    "declara_control": False,          # <- no declara: debe salir la carta
    "es_empresa_familiar": True,
    "nucleo_familiar": [
        {"tipo_doc": "C.C.", "id_num": "71111111", "nombre": "Juan Pablo García",
         "acciones": "12.000", "parentesco": "Padre"},
        {"tipo_doc": "C.C.", "id_num": "43222222", "nombre": "María Camila Torres",
         "acciones": "8.000", "parentesco": "Hija"},
    ],
    "apoderado": None,
}


def test_escenario_a(client):
    archivos = _generar(client, _payload("payload_a.json", PAYLOAD_A), "escenario_a")
    nombres = sorted(archivos)
    print("\n  Escenario A — documentos generados:")
    for n in nombres:
        print("    ·", n)

    # ── Documentos esperados ──
    assert _buscar(archivos, "Estatutos.docx")
    assert _buscar(archivos, "Formulario_RUES")
    assert _buscar(archivos, "Otras_Entidades")
    assert _buscar(archivos, "Responsabilidades_Tributarias")
    assert _buscar(archivos, "Emprendimiento_Social")
    assert _buscar(archivos, "Grupo_Etnico")
    assert _buscar(archivos, "Empresa_Familiar"), "falta el formato de empresa familiar"
    assert _buscar(archivos, "Carta_No_Situacion_Control"), "falta la carta de no control"
    assert not _buscar(archivos, "Formato_Situacion_Control"), \
        "no debía generarse el formato de situación de control"

    # ── Estatutos ──
    est = _docx_txt(_buscar(archivos, "Estatutos.docx"))
    assert "{{" not in est
    assert "QUINIENTOS MILLONES DE PESOS (COP $500.000.000)" in est
    assert "cinco millones (5.000.000) acciones" in est
    assert "cien pesos moneda legal colombiana (COP $100)" in est
    assert "veinte mil (20.000) acciones" in est          # suscrito 2.000.000/100
    assert "Junta directiva:" in est
    assert "tres (3) miembros principales" in est
    assert "PETER SCHMIDT, identificado con Pasaporte No. X8899221" in est
    assert "LAURA GIL PEÑA, identificada con C.C. No. 43555444" in est
    assert "Revisor fiscal:" in est
    assert "AUDITORES ASOCIADOS S.A.S., sociedad identificada con NIT 900.111.222-3" in est
    assert "CARLOS MESA URIBE" in est
    assert "tarjeta profesional No. 12345-T" in est
    # El tipo de documento se escribe igual en todos los nombramientos: el
    # formulario envía "CC" y en los estatutos debe leerse "C.C."
    assert "Representante legal principal: JUAN PABLO GARCÍA, identificado con C.C. No." in est
    assert " con CC No. " not in est, "quedó un tipo de documento sin normalizar"

    # ── Carta de no control ──
    carta = _pdf_txt(_buscar(archivos, "Carta_No_Situacion_Control"))
    assert "FAMILIA ANDINA S.A.S." in carta
    assert "accionista mayoritario, JUAN PABLO GARCÍA, titular del 60% de las acciones suscritas" in carta
    assert "artículo 261 del Código de Comercio" in carta
    assert "C.C. No. 71111111" in carta
    assert "CC No." not in carta, "quedó un tipo de documento sin normalizar en la carta"
    # La cámara es la elegida en el Paso 1, no la del municipio
    assert "CÁMARA DE COMERCIO DE ABURRÁ SUR" in carta, \
        "la carta debe dirigirse a la cámara elegida, no a la del municipio"

    # ── Empresa familiar ──
    fam = _pdf_txt(_buscar(archivos, "Empresa_Familiar"))
    assert "LEY 2495 DE 2025" in fam
    assert "FAMILIA ANDINA S.A.S." in fam
    assert "Padre" in fam and "Hija" in fam
    assert "12.000" in fam and "8.000" in fam
    assert "ABURRÁ SUR" in fam, "el formato debe llevar la cámara elegida"
    assert "Envigado" in fam
    assert "C.C. 71111111" in fam, "el tipo de documento del RL debe ir normalizado"
    # El formato se estampa con reportlab, no en AcroForm: conserva las tildes
    assert "JUAN PABLO GARCÍA" in fam, "los nombres propios deben conservar tildes"
    assert "María Camila Torres" in fam

    print("  OK  escenario A")


# ════════════════════════════════════════════════════════════════
# ESCENARIO B
# ════════════════════════════════════════════════════════════════
PAYLOAD_B = {
    "nombre_sas": "SOLO UNO S.A.S.",
    "municipio": "Medellín", "departamento": "Antioquia",
    "direccion": "Calle 10 #40-20", "barrio": "Manila",
    "email": "hola@solouno.co", "telefono1": "3009998877",
    "telefono2": "", "telefono3": "",
    "zona": "urbana", "tipo_local": "oficina", "tenencia": "propia",
    "accionistas": [
        {"tipo": "natural", "nombre": "Ana Restrepo Gómez", "tipo_doc": "CC",
         "id_tipo": "C.C.", "id_num": "43000111", "expedicion": "Medellín",
         "domicilio": "Medellín", "nacimiento": "", "genero": "F",
         "porcentaje": 100, "capital_pagado": ""},
    ],
    "rl_principal": {"nombre": "Ana Restrepo Gómez", "cc": "43000111",
                     "tipo_doc": "C.C.", "expedicion": "Medellín", "genero": "F"},
    "rl_suplente": None,
    "ciiu_code": "7020", "ciiu_description": "Actividades de consultoría de gestión",
    "ciiu_code_sec": "", "ciiu_description_sec": "",
    "objeto_social": "Consultoría empresarial.",
    "capital_autorizado": "1.000.000.000",
    "valor_nominal": "1",
    "capital_suscrito": "1.000.000",
    "capital_pagado": "",
    "regimen": "ordinario", "ingresos_mensuales": "100.000",
    "tiene_junta": False, "tiene_revisor": False,
    "junta_directiva": None, "revisor_fiscal": None,
    "es_emprendimiento_social": False, "grupo_etnico": "",
    "declara_control": True,           # <- sí declara: formato oficial
    "es_empresa_familiar": False, "nucleo_familiar": [],
    "apoderado": None,
}


def test_escenario_b(client):
    archivos = _generar(client, _payload("payload_b.json", PAYLOAD_B), "escenario_b")
    print("\n  Escenario B — documentos generados:")
    for n in sorted(archivos):
        print("    ·", n)

    assert _buscar(archivos, "Formato_Situacion_Control"), "falta el formato de situación de control"
    assert not _buscar(archivos, "Carta_No_Situacion_Control")
    assert not _buscar(archivos, "Empresa_Familiar")

    est = _docx_txt(_buscar(archivos, "Estatutos.docx"))
    assert "{{" not in est
    assert "un peso moneda legal colombiana (COP $1)" in est
    assert "Junta directiva:" not in est
    assert "Revisor fiscal:" not in est
    assert "identificada con C.C. No. 43000111" in est   # concordancia femenina
    print("  OK  escenario B")


# ════════════════════════════════════════════════════════════════
# VALIDACIONES DEL BACKEND
# ════════════════════════════════════════════════════════════════
def test_validaciones(client):
    malo = dict(PAYLOAD_B, capital_autorizado="500.000", capital_suscrito="1.000.000")
    r = client.post("/api/generate", json=malo)
    assert r.status_code == 400, "debía rechazar autorizado < suscrito"
    assert "autorizado" in r.get_json()["error"].lower()

    malo2 = dict(PAYLOAD_B, valor_nominal="300", capital_suscrito="1.000.000")
    r2 = client.post("/api/generate", json=malo2)
    assert r2.status_code == 400, "debía rechazar capital no múltiplo del nominal"
    assert "múltiplos" in r2.get_json()["error"]
    print("  OK  validaciones de capital")


if __name__ == "__main__":
    cli = _cliente()
    test_escenario_a(cli)
    test_escenario_b(cli)
    test_validaciones(cli)
    print(f"\nTodas las pruebas del paquete pasaron. Salida en: {OUT}")
